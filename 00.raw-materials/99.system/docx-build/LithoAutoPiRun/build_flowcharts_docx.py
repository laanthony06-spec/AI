from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_requirement_docx import (
    BLACK,
    MUTED,
    add_field,
    add_figure,
    set_run_font,
    setup_page,
    setup_styles,
)
from make_bw_flowcharts import make_flowcharts


ROOT = Path(r"D:\Obsidian\work\OBSidianCodex")
BUILD_DIR = ROOT / "00.raw-materials" / "99.system" / "docx-build" / "LithoAutoPiRun"
ASSET_DIR = BUILD_DIR / "assets"
OUT_DIR = ROOT / "00.raw-materials" / "90.processed" / "LithoAutoPiRun"
FLOW_DIR = OUT_DIR / "流程图"
OUT_PATH = OUT_DIR / "LithoAutoPiRun_复杂逻辑流程图.docx"
QA_DIR = BUILD_DIR / "flowcharts-qa"
QA_PATH = QA_DIR / "flowcharts.docx"


FLOWCHARTS = [
    (
        "01_RTD候选筛选与Pilot选择流程.png",
        "01-rtd-selection.png",
        "流程图 1：RTD 候选筛选与 Pilot 选择",
        "展示 Lot 获取、基础条件过滤、候选排序与 Pilot 选定的完整判断链路。",
        6.15,
    ),
    (
        "02_Pilot整批与物理分批判断流程.png",
        "02-pilot-split.png",
        "流程图 2：Pilot 整批与物理分批判断",
        "展示整批条件、分批数量配置校验、Wafer 排序、FOUP 预留及失败回退逻辑。",
        6.15,
    ),
    (
        "03_WaitPilotChangeFOUP卡控流程.png",
        "03-wait-pilot-control.png",
        "流程图 3：WaitPilotChangeFOUP 卡控",
        "展示 Litho、BARCO 与其他站点在不同 Queue Time 状态下的卡控差异。",
        6.15,
    ),
    (
        "04_AMA物理分批与回退流程.png",
        "04-ama-split.png",
        "流程图 4：AMA 物理分批与回退",
        "展示分批前置校验、Transfer FOUP、物理分批执行与异常回退处理。",
        6.15,
    ),
]


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = p.add_run("LithoAutoPiRun 复杂逻辑流程图")
    set_run_font(r, size=8.5, color=BLACK, bold=True)
    r = p.add_run("\t独立交付版 | 2026-07-24")
    set_run_font(r, size=8.5, color=BLACK)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=BLACK)
    add_field(p, "PAGE", "1")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=BLACK)


def add_page_title(doc, title, summary):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=17, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(summary)
    set_run_font(r, size=9.5, color=BLACK)


def build_document():
    make_flowcharts()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FLOW_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_styles(doc)
    section = setup_page(doc)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    setup_header_footer(section)

    for index, (public_name, asset_name, title, summary, width) in enumerate(FLOWCHARTS):
        source = ASSET_DIR / asset_name
        public_path = FLOW_DIR / public_name
        shutil.copy2(source, public_path)

        if index:
            doc.add_page_break()
        add_page_title(doc, title, summary)
        add_figure(
            doc,
            "复杂逻辑流程图（按最终确认规则整理）",
            source,
            f"{title}。{summary}",
            width_inches=width,
        )

    doc.save(OUT_PATH)
    shutil.copy2(OUT_PATH, QA_PATH)
    print(OUT_PATH)
    for public_name, *_ in FLOWCHARTS:
        print(FLOW_DIR / public_name)


if __name__ == "__main__":
    build_document()
