from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_requirement_docx import (
    BLACK,
    BLUE,
    INK_BLUE,
    LIGHT_GRAY,
    MUTED,
    add_approval_table,
    add_body,
    add_bullets,
    add_field,
    add_figure,
    add_heading,
    add_numbers,
    add_rich_text,
    set_cant_split,
    set_cell_shading,
    set_run_font,
    set_table_borders,
    set_table_geometry,
    setup_page,
    setup_styles,
)
from make_bw_flowcharts import make_flowcharts


ROOT = Path(r"D:\Obsidian\work\OBSidianCodex")
BUILD_DIR = ROOT / "00.raw-materials" / "99.system" / "docx-build" / "LithoAutoPiRun"
ASSET_DIR = BUILD_DIR / "assets"
OUT_DIR = ROOT / "00.raw-materials" / "90.processed" / "LithoAutoPiRun"
OUT_PATH = OUT_DIR / "LithoAutoPiRun_需求申请单_最终版.docx"
QA_DIR = BUILD_DIR / "final-qa"
QA_PATH = QA_DIR / "final-requirement.docx"


def force_all_text_black(doc):
    """Keep every visible Word text element black, including headers and tables."""
    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.color.rgb = RGBColor(0, 0, 0)

    def blacken(container):
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, color="000000")
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    blacken(cell)

    blacken(doc)
    for section in doc.sections:
        blacken(section.header)
        blacken(section.footer)


def apply_body_typography(doc):
    """Apply the requested正文字体 and remove all character-level shading."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)

    def format_run(run, body=False):
        rpr = run._element.get_or_add_rPr()
        for shading in list(rpr.findall(qn("w:shd"))):
            rpr.remove(shading)
        if body:
            set_run_font(
                run,
                ascii_font="Times New Roman",
                east_asia="SimSun",
                size=10.5,
                color="000000",
            )

    body_started = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "一、基础信息":
            body_started = True
        for run in paragraph.runs:
            format_run(run, body=body_started and paragraph.style.name == "Normal")

    def format_tables(container):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            format_run(run, body=True)
                    format_tables(cell)

    format_tables(doc)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    format_run(run, body=False)
            format_tables(container)


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = p.add_run("LithoAutoPiRun 功能需求规格")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = p.add_run("\t最终版 | 2026-07-25")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(p, "PAGE", "1")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_metadata_table(doc):
    rows = [
        ("编号", "由信息技术部填写", "类别", "功能开发"),
        ("申请部门", "制造部", "申请人员", "温浩奇"),
        ("系统名称", "CIM 计算机集成制造系统 Fab6（一科）", "功能模块", "智能派工系统（RTD／DSP）"),
        ("申请日期", "2026-07-25", "希望交付日期", "2026-07-29"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1400, 3280, 1400, 3280])
    set_table_borders(table)
    for row, values in zip(table.rows, rows):
        set_cant_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, value, size=9.7, bold=(idx % 2 == 0))
            if idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flowchart(doc, title, image_name, caption, first=False):
    heading = add_heading(doc, title, 2)
    if not first:
        heading.paragraph_format.page_break_before = True
    add_figure(doc, caption, ASSET_DIR / image_name, caption, width_inches=5.9)


def build_document():
    make_flowcharts()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_styles(doc)
    section = setup_page(doc)
    setup_header_footer(section)

    # memo_masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("功能需求规格")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("LithoAutoPiRun 需求申请单")
    set_run_font(r, size=24, color=INK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("逻辑分批调整为物理分批，并优化 Pilot 选择、Transfer FOUP 及相关卡控")
    set_run_font(r, size=12.5, color=MUTED)

    add_heading(doc, "一、基础信息", 1)
    add_metadata_table(doc)

    add_heading(doc, "二、项目简介和必要性分析", 1)
    add_heading(doc, "2.1 项目背景", 2)
    add_body(doc, "当前 `LithoAutoSplitPirun` 采用逻辑分批。生成的 Pilot 与母批仍位于同一 FOUP，跨厂转运时可能受同 FOUP 内其他 Lot 状态限制，Pilot 无法及时执行 PiRun。")

    add_heading(doc, "2.2 存在问题", 2)
    add_bullets(doc, [
        "逻辑分批未实际拆分 Wafer，Pilot 与母批共用 FOUP，Transfer 条件受同载具其他 Lot 影响。",
        "现有 AutoPiRun 只对满足分批条件的 Lot 自动设置 Pilot；不满足条件的 Lot 可能持续卡控。",
        "Pilot 选择、物理分批、Transfer FOUP、子母批同机台卡控分散在 RTD、Rule、LithoAssign 和 AMA，缺少统一闭环。",
    ])

    add_heading(doc, "2.3 需求目标", 2)
    add_bullets(doc, [
        "将 `LithoAutoSplitPirun` 由逻辑分批调整为物理分批。",
        "优化 RTD 候选 Lot 筛选、Context 排序、Pilot 选片及 Report 输出。",
        "由 AMA 完成分批前复核、空 FOUP 预占、MES 物理分批和 Transfer FOUP。",
        "补充 `WaitPilotChangeFOUP`、`R2RAutoPirunControl`、`Parent&ChildLotNeedRunSameTool` 卡控。",
    ])

    add_heading(doc, "2.4 本次范围", 2)
    add_body(doc, "**涉及范围：**RTD Report、Global Macro、LithoRule、LithoAssign、AMA、MES 物理分批与 Transfer FOUP。")
    add_body(doc, "**不在本次范围：**不新增重复物理分批防护；不新增 `request_id`、`generated_time`、`execution_status`、`error_code`；不重新定义 WatchDog 时区或 Trigger Time Slot 边界；详细 Test Case 另行编制。")

    add_heading(doc, "三、项目投资方案比较及效果分析", 1)
    add_heading(doc, "3.1 原方案", 2)
    add_body(doc, "维持逻辑分批。该方案改动较小，但 Pilot 与母批仍共用 FOUP，无法解决跨厂及同 FOUP Other Lot 对 Pilot 转运的限制。")

    add_heading(doc, "3.2 改善方案", 2)
    add_body(doc, "RTD 选择 Pilot Wafer 并输出 `IsNeedSplit` 与 `pi_splitwafer`；AMA 复核 Lot 状态后执行物理分批，必要时为 Pilot 更换 FOUP，并同步补充相关 Rule 卡控。")

    add_heading(doc, "3.3 效果分析", 2)
    add_bullets(doc, [
        "Pilot 与母批可分离处理，降低同 FOUP Other Lot 对 PiRun 执行的影响。",
        "无空 FOUP、MES 分批失败等场景有明确回退路径。",
        "降低 Lot Queue Time 超时风险。",
    ])

    flow_heading = add_heading(doc, "四、需求流程图", 1)
    flow_heading.paragraph_format.page_break_before = True
    add_body(doc, "复杂逻辑拆分为四个子流程，流程节点与正文判断规则一致。")
    add_flowchart(doc, "4.1 RTD 候选筛选与 Pilot 选择", "01-rtd-selection.png",
                  "图 1  RTD 候选筛选与 Pilot 选择流程", first=True)
    add_flowchart(doc, "4.2 Pilot 整批与物理分批判断", "02-pilot-split.png",
                  "图 2  Pilot 整批与物理分批判断流程")
    add_flowchart(doc, "4.3 WaitPilotChangeFOUP 卡控", "03-wait-pilot-control.png",
                  "图 3  WaitPilotChangeFOUP 卡控流程")
    add_flowchart(doc, "4.4 AMA 物理分批与回退", "04-ama-split.png",
                  "图 4  AMA 物理分批与回退流程")

    rtd_heading = add_heading(doc, "五、RTD 部分更改逻辑", 1)
    rtd_heading.paragraph_format.page_break_before = True
    add_body(doc, "RTD 修改范围包括两个 Report、Global Macro、LithoRule 和 LithoAssign。`Pi_split_flag` 统一使用 `Y` 表示允许物理分批；PiRun 状态统一使用 `PIRUNON`。")

    add_heading(doc, "5.1 数据获取与基础筛选", 2)
    add_heading(doc, "5.1.1 数据范围", 3)
    add_body(doc, "从 FAB6、FAB8 获取候选 Lot。读取以下资料：")
    add_bullets(doc, [
        "`fwlot`：`appid`、`priority`、`processingstatus`、`componentqty`。",
        "`fablotext`：`requiredcapability`、`runcardid`、`reticleid`。",
        "`fablotcarrierext`：`carrierkind`。",
        "`fabinqtimeprocess`：`RemainQ`。",
        "`UI.RTDConfig-LITHOLotAssignment-LithoAssignCapability`：`LithoCapability`。",
    ])
    add_body(doc, "BARCO Capability 固定为 `L-BARCO-L`、`L-BARCO-S`。")

    add_heading(doc, "5.1.2 基础筛选", 3)
    add_body(doc, "候选 Lot 必须同时满足以下条件：")
    add_numbers(doc, [
        "`processingstatus` 为 `Active` 或 `CrossFabTransferred`。",
        "`carrierkind='FOUP'`。",
        "`runcardid` 为空。",
        "`requiredcapability` 为 `LithoCapability`、`L-BARCO-L` 或 `L-BARCO-S`。",
    ])
    add_body(doc, "跨厂 Lot 沿用现有 `istransferlot` Macro。若 `IsTransferLot=True`，保留该记录；若 `IsTransferLot<>True`，则仅保留 `processingstatus<>'CrossFabTransferred'` 的记录。本需求不新增 FAB6/FAB8 跨厂资料去重规则。")

    add_heading(doc, "5.2 PiRun 站点与 R2R 条件", 2)
    add_heading(doc, "5.2.1 PiRunLoop", 3)
    add_body(doc, "对通过基础筛选的 Lot 向后 Fetch 20 个站点，读取 `productname`、`planname`、`stage`、`capability`、`stepseq`、`recipeid`、`STN`。Transfer Lot 的厂别信息沿用现有修正规则。")
    add_body(doc, "Lot 当前站点至相同 Stage 最后一道 CD 站点之间的区段定义为 `PiRunLoop`。若 Fetch 结果不存在 CD 站点，则剔除该 Lot。")

    add_heading(doc, "5.2.2 进一步筛选", 3)
    add_numbers(doc, [
        "`PiRunLoop` 中必须存在 `capability=LithoCapability` 且具有 Reticle 信息的 Litho 站点。",
        "按 `productid + layer + lotid` 匹配 `r2r_litho_whitelist`；命中时沿用 Specify Lot 处理。",
        "同一 FOUP 已存在 Litho Pilot 时，剔除当前候选 Lot。",
        "`PiRunLoop` 中存在 FutureHold 或 RC 站点时，剔除该 Lot。",
    ])

    add_heading(doc, "5.2.3 可作业机台与 Context", 3)
    add_body(doc, "通过现有 Transfer Macro 判断 Lot 是否可在对应 Litho 站点执行 Transfer；满足时获取对应工厂的 Litho 机台。按厂别从 `rtd_r2r_litho_add_setting` 读取 `Pi_split_flag` 和 `pi_splitcnt`。若不存在 `Pi_split_flag='Y'` 的可作业机台，则剔除该 Lot。")
    add_body(doc, "继续检查 EQPStatus、LCC、Capability、Recipe、PPID 和 Global Reason；存在卡控时剔除对应机台。")
    add_body(doc, "按 `Lot + STN + Reticle` 从 `rtd_r2r_litho_context_ovl`、`rtd_r2r_litho_context_cd` 匹配 Context。OVL、CD 状态均须属于 `PIRUNON`、`ON`、`Fixed`，且不存在 R2R Reason。若同一 Lot 的有效 `ContextCount>1`，按多路径 Lot 剔除。")

    add_heading(doc, "5.3 Pilot 候选排序与循环选择", 2)
    add_heading(doc, "5.3.1 Context 内 Lot 排序指标", 3)
    add_bullets(doc, [
        "`GapToLitho`：Lot 当前站点距 Litho 站点的剩余 Step 数。",
        "`SplitCntMatched`：`componentqty>=pi_splitcnt` 时为 1，否则为 0。`pi_splitcnt` 无效时按默认值 4 计算。",
        "`RequiredChuckCount`：存在 `prelayer` 时读取 `r2r_litho_waferhistory` 的 Chuck；否则读取 `fsmaterialassociation` 的 Slot。C1/C2 或奇偶 Slot 各至少 2 片时为 1，否则为 0。",
        "`BulletLot`：沿用现有空扣或空 LP Lot 判断。",
        "`RemainQ`：取实际剩余 Queue Time；空值按 `9999 h`。",
        "`KeyLot`：`quota_applyinfo` 中 `KeyLot=1` 且 `Status='CONFIRM'` 时为 1，否则为 0。",
    ])
    add_body(doc, "`RTDRank` 排序优先级如下：")
    add_numbers(doc, [
        "`GapToLitho ASC`。",
        "`SplitCntMatched DESC`。",
        "`RequiredChuckCount DESC`。",
        "`BulletLot DESC`。",
        "`RemainQ ASC`。",
        "`KeyLot ASC`，即普通 Lot 优先于 Key Lot。",
        "`lotid ASC`，作为最终 Tie-breaker。",
    ])

    add_heading(doc, "5.3.2 Context 排序与循环", 3)
    add_bullets(doc, [
        "`ReticleSTNRank`：第一轮优先 Reticle 当前所在 STN；后续轮次优先与上一轮相同的 `Reticle + STN` 分组。",
        "`ContextCandidateCount`：当前 Context 的候选 Lot 数量。",
        "`ActualSTNPilotCount`：对应 STN 已选 Pilot Context 数量。",
    ])
    add_body(doc, "最终按 `ReticleSTNRank DESC`、`ContextCandidateCount ASC`、`ActualSTNPilotCount ASC`、`RTDRank ASC` 排序。")
    add_body(doc, "每轮固定排序第一的 `Lot + Context`。每个 Context 最多选择一个 Pilot；选中后删除同一 Context 的其他候选，并重新计算排序指标，直至无可用 Context 或无可用 Lot。")

    add_heading(doc, "5.4 Pilot 整批与物理分批判断", 2)
    add_heading(doc, "5.4.1 整批 Pilot 条件", 3)
    add_body(doc, "以下五项使用“或”关系。任意一项成立时，设置 `IsNeedSplit=F`，整批 Lot 作为 Pilot：")
    add_numbers(doc, [
        "`BulletLot=1` 或 `KeyLot=1`。",
        "`CurCapability=LithoCapability` 且 `FuLL(RemainQ)`。",
        "`RequiredChuckCount=0`。",
        "`SplitCntMatched=0`。",
        "`componentqty<=6`。",
    ])
    add_body(doc, "`componentqty>0` 由上游保证，本需求不增加零片判断。")

    add_heading(doc, "5.4.2 物理分批选片", 3)
    add_body(doc, "不满足任一整批条件时，设置 `IsNeedSplit=T`。Wafer ID 1～10 为 Group1，11～25 为 Group2；有 Chuck 信息时 C1/C2 对应 SubGroup1/SubGroup2，无 Chuck 信息时奇数/偶数 Slot 对应 SubGroup1/SubGroup2。")
    add_bullets(doc, [
        "Group1 + SubGroup1：`GroupRank=1`。",
        "Group1 + SubGroup2：`GroupRank=2`。",
        "Group2 + SubGroup1：`GroupRank=3`。",
        "Group2 + SubGroup2：`GroupRank=4`。",
    ])
    add_body(doc, "在每个 `Group + SubGroup` 内按 Wafer ID 升序计算 `WaferRank`，最终按 `WaferRank ASC`、`GroupRank ASC` 轮流选片。")
    add_bullets(doc, [
        "`pi_splitcnt` 为空、等于 0、小于 0 或大于 25 时，使用默认值 4。",
        "`pi_splitcnt` 在 1～25 范围内时，按配置值选片。",
        "选片数大于当前可用 Wafer 数时，不执行物理分批，改为整批 Pilot。",
        "选中 Wafer 写入 `pi_splitwafer`。",
    ])

    add_heading(doc, "5.4.3 Merge 站点", 3)
    add_body(doc, "若 Lot 存在不包含 SRC 的 ADI 站点，则取第一道符合条件的 ADI 作为 Merge 站点；否则取最后一道 CD。有效 Merge 站点由现有流程保证，本需求不增加两者均不存在时的处理。")

    add_heading(doc, "5.5 Report 调整", 2)
    add_heading(doc, "5.5.1 Central_GetLithoR2RAutoPirunInfo", 3)
    add_body(doc, "修改现有 Pilot 选择和选片逻辑，输出字段包括：")
    add_bullets(doc, [
        "Lot、`toolid`、`productid`、`layerid`、`reticleid`。",
        "`prereticle`、`pretool`、`custom_context_value`。",
        "`pi_splitwafer`、`IsNeedSplit`、`isSTNSite`。",
    ])
    add_body(doc, "保持现有输出范围和记录方式，不新增请求追踪或执行状态字段。")

    add_heading(doc, "5.5.2 LithoPilotAutoDoAdhocSorter", 3)
    add_body(doc, "新增 Report，用于取得需要 Transfer FOUP 的 Litho Pilot。")
    add_bullets(doc, [
        "从 `r2r_litho_context_ovl` 读取 `ovl_status`、`pilot`。",
        "从 `r2r_litho_context_cd` 读取 `cd_status`、`pilot`。",
        "从 `fwlot` 读取 `appid`、`lottype`、`priority`。",
        "从 `fabcategorymap` 读取 `lottype`、`category`。",
    ])
    add_body(doc, "保留 `priority<5`、`category='Production'`，且 OVL 或 CD 状态为 `PIRUNON` 的记录。按 Carrier 获取同 FOUP 所有 Lot，并保留 `extrastatus='WaitForJobPrep'` 的记录；若 Carrier 中存在 Other Lot，则该 Pilot 需要 Change FOUP。")
    add_body(doc, "Carrier 按 `RemainQ ASC`、`Priority ASC`、`componentqty DESC`、`lotid ASC` 排序并建立 `AdhocSorterJob`。")
    add_body(doc, "当 `AMA.TriggerConfig-WatchDog_LithoPilotAutoDoAdhocSorter.Switch='Y'` 且当前时间位于现有 Trigger Time Slot 时，按 `TriggerCount/Time` 选择 Carrier。输出字段包括 Carrier、Pilot、`extrastatus`、Status、RemainQ、Pieces、Prod、Priority。")

    add_heading(doc, "5.6 Rule 与 LithoAssign", 2)
    add_heading(doc, "5.6.1 Global Macro：WaitPilotChangeFOUP", 3)
    add_body(doc, "WatchDog 开启并位于 Trigger Time Slot 时，按 Pilot 是否位于 `UnscheduledSorter` 决定需要卡控的同 FOUP Other Lot。符合范围的 Other Lot 增加 `Reason=WaitPilotChangeFOUP`。")
    add_numbers(doc, [
        "Other Lot 位于 Litho 站点：不解除卡控。",
        "Other Lot 位于 BARCO 站点：`RemainQ<4 h` 或触发 `Qu_0` 时解除。",
        "Other Lot 位于其他站点：仅触发 `Qu_0` 时解除。",
        "Pilot 的 `RemainQ<4 h` 或 Pilot 触发 `Qu_0` 时，解除相关卡控。",
    ])

    add_heading(doc, "5.6.2 LithoRule：R2RAutoPirunControl", 3)
    add_body(doc, "针对非 Specify Lot，若 `Pi_split_flag='Y'`，R2R CD 或 OVL Status 为 `PIRUNON`，且 Pilot 不为空，则增加 `Reason=R2RAutoPirunControl`；否则不增加。")

    add_heading(doc, "5.6.3 LithoRule：Parent&ChildLotNeedRunSameTool", 3)
    add_body(doc, "从 `r2r_litho_context_relation` 获取 `productid`、`curr_layer`、`pre_layer`，从 `r2r_litho_context_ovl` 获取 `productid`、`layerid`、`pretool`。只有 `pretool` 不为空时继续判断。")
    add_body(doc, "通过 `fabfutureaction` 取得 FutureMerge 关系的子批或母批；关系范围沿用现有返回结果。从 `r2r_lot_history` 按 `lotid + productid + layerid` 获取对应 Layer 的 `toolid`，以实际作业完成时间降序、记录 ID 降序取得最新记录。")
    add_body(doc, "针对非 Specify Lot，若当前 Lot 与子批或母批的 `toolid` 不一致，则增加 `Reason=Parent&ChildLotNeedRunSameTool`；否则继续现有逻辑。")

    add_heading(doc, "5.6.4 LithoAssign", 3)
    add_body(doc, "LithoAssign 保留 `R2RAutoPirunControl`，并新增 `Parent&ChildLotNeedRunSameTool`。判断逻辑与 LithoRule 一致，`pretool` 及子母批作业机台信息由 Central 获取。")

    add_heading(doc, "5.7 UI 配置", 2)
    add_bullets(doc, [
        "`rtd_r2r_litho_add_setting.Pi_split_flag`：`Y` 表示机台允许物理分批。",
        "`rtd_r2r_litho_add_setting.pi_splitcnt`：有效范围 1～25；为空、0、负数或大于 25 时使用默认值 4。",
        "`AMA.TriggerConfig-WatchDog_LithoPilotAutoDoAdhocSorter`：沿用 `Switch`、`Trigger Time Slot`、`TriggerCount/Time`。",
    ])

    ama_heading = add_heading(doc, "六、AMA 部分更改逻辑", 1)
    ama_heading.paragraph_format.page_break_before = False
    add_heading(doc, "6.1 触发方式与数据获取", 2)
    add_body(doc, "Pilot 设置沿用现有 AMA 触发方式和执行周期，本需求不新增 Trigger。AMA 读取 `Central_GetLithoR2RAutoPirunInfo` 执行 Pilot 设置；Transfer FOUP 按 WatchDog 配置读取 `LithoPilotAutoDoAdhocSorter`。")
    add_body(doc, "设置 Pilot 时读取 Lot、`toolid`、`productid`、`layerid`、`reticleid`、`prereticle`、`pretool`、`custom_context_value`、`pi_splitwafer`、`IsNeedSplit`、`isSTNSite`。")

    add_heading(doc, "6.2 整批 Pilot", 2)
    add_body(doc, "若 `IsNeedSplit=F`，则直接将整批 Lot 传给 R2R。")

    add_heading(doc, "6.3 物理分批 Pilot", 2)
    add_heading(doc, "6.3.1 执行前复核", 3)
    add_body(doc, "若 `IsNeedSplit=T`，AMA 必须重新检查以下六项：")
    add_numbers(doc, [
        "Lot 属于当前执行工厂，即 FAB6 或 FAB8。",
        "`fwlot.extrastatus='WaitForJobPrep'`。",
        "Lot 当前站点的 `runcardid` 为空。",
        "当前 Capability 为 `LithoCapability`、`L-BARCO-L` 或 `L-BARCO-S`。",
        "`CarrierKind='FOUP'`。",
        "Report 选中的 Wafer 当前仍归属于该 Lot；不额外检查 Wafer 状态、Slot 或 Chuck。",
    ])
    add_body(doc, "若任一项不满足，则停止处理该 Lot，记录失败原因，等待下一轮重新计算；不回退为整批 Pilot。现有系统负责避免同一 Lot 重复物理分批，本需求不新增幂等防护。")

    add_heading(doc, "6.3.2 空 FOUP 与 MES 分批", 3)
    add_body(doc, "六项复核通过后，先获取并预占空 FOUP，再调用 MES 物理分批接口。")
    add_bullets(doc, [
        "未取得空 FOUP：不执行物理分批，将整批 Lot 传给 R2R。",
        "接口成功：将 `pi_splitwafer` 分出，生成子批并设置为 Pilot，再按现有顺序传给 R2R、执行 Transfer FOUP。",
        "接口失败：立即释放预占 FOUP，将整批 Lot 作为 Pilot 传给 R2R；`IsNeedSplit`、`pi_splitwafer` 的后续处理沿用现有 AMA 逻辑。",
    ])

    add_heading(doc, "6.4 Transfer FOUP", 2)
    add_body(doc, "AMA 按 `LithoPilotAutoDoAdhocSorter` 排序结果调用 MES Transfer FOUP 接口，将 Pilot 导入已取得的空 FOUP。")
    add_body(doc, "若无法取得空 FOUP、可用空 FOUP 数量为 0，或 MES Transfer FOUP 接口失败，则在 `AMALog` 记录失败信息。Other Lot 紧急解除卡控后，Pilot 的 Transfer FOUP 任务继续由现有 Adhoc Sorter 流程处理。")

    add_heading(doc, "6.5 AMA 输出", 2)
    add_bullets(doc, [
        "整批 Pilot 或物理分批后的子批 Pilot 传给 R2R。",
        "需要 Change FOUP 的 Pilot 按 Report 顺序执行 Transfer FOUP。",
        "前置复核、物理分批或 Transfer FOUP 失败时记录原因。",
        "本需求不新增 Alarm。",
    ])

    add_heading(doc, "七、原逻辑／修改后逻辑／修改原因", 1)
    add_heading(doc, "7.1 原逻辑", 2)
    add_bullets(doc, [
        "`LithoAutoSplitPirun` 使用逻辑分批，Pilot 与母批仍位于同一 FOUP。",
        "AutoPiRun 主要对满足既有条件的 Lot 设置 Pilot，Transfer FOUP 与物理分批未形成完整闭环。",
        "子母批同机台、WaitPilotChangeFOUP 等卡控未按本需求规则处理。",
    ])

    add_heading(doc, "7.2 修改后逻辑", 2)
    add_bullets(doc, [
        "RTD 统一筛选 FAB6/FAB8 候选 Lot，按 Context 和 RTDRank 循环选择 Pilot。",
        "RTD 输出整批或物理分批结果；AMA 复核后执行物理分批，并处理空 FOUP 和接口失败回退。",
        "新增 Change FOUP Report 及 WaitPilotChangeFOUP、R2RAutoPirunControl、Parent&ChildLotNeedRunSameTool 卡控。",
    ])

    add_heading(doc, "7.3 修改原因", 2)
    add_body(doc, "逻辑分批无法解除 Pilot 与母批共用 FOUP 的限制。改为物理分批并增加 Transfer FOUP 后，Pilot 可独立进入 PiRun 流程；同时通过 Rule 和 AMA 复核避免状态变化导致错误执行。")

    add_heading(doc, "八、影响范围及异常边界", 1)
    add_heading(doc, "8.1 系统影响", 2)
    add_bullets(doc, [
        "RTD：影响候选 Lot Filter、Context Sorting、Pilot 选片、Reason、Report 和 LithoAssign。",
        "AMA：影响 Pilot 设置、物理分批、空 FOUP 预占、Transfer FOUP 和失败记录。",
        "MES：调用现有物理分批与 Transfer FOUP 接口；接口成功或失败按本需求分支处理。",
        "MCS／AMHS：本需求不新增搬送规则；MES 后续搬送沿用现有流程。",
        "EAP：本需求不修改 EAP 逻辑。",
        "DB：本需求未定义新增业务表或字段；两个 Report 的输出字段按 5.5 执行。",
    ])

    add_heading(doc, "8.2 边界与异常", 2)
    add_bullets(doc, [
        "`RemainQ` 为空时按 `9999 h` 参与排序。",
        "`pi_splitcnt` 无效时使用默认值 4；有效但大于当前可用 Wafer 数时改为整批 Pilot。",
        "候选 Lot 无 CD、无有效 Litho/Reticle、存在 FutureHold/RC、无有效机台或多路径时直接剔除。",
        "AMA 六项复核失败时停止并等待下一轮，不回退整批。",
        "无空 FOUP 时整批回退；MES 物理分批失败时释放 FOUP 后整批回退。",
        "MES Transfer FOUP 失败时记录 `AMALog`，不新增额外恢复动作。",
        "WatchDog 时间判断、Specify Lot、Transfer Macro、FuLL(RemainQ)、Adhoc Sorter 后续处理均沿用现有逻辑。",
    ])

    add_heading(doc, "九、待确认事项", 1)
    add_numbers(doc, [
        "需求编号由信息技术部填写。",
        "新增 Report `LithoPilotAutoDoAdhocSorter` 的唯一记录维度、同一次执行的去重方式及历史保留策略需 IT 确认；本需求不新增对应业务字段。",
    ])

    add_heading(doc, "十、测试及验收说明", 1)
    add_body(doc, "详细 Test Case 另行建立。本需求验收至少覆盖以下场景：")
    add_bullets(doc, [
        "整批 Pilot 与物理分批 Pilot。",
        "`pi_splitcnt` 空值、0、负数、大于 25、以及大于可用 Wafer 数。",
        "FAB6／FAB8 与 `CrossFabTransferred`。",
        "AMA 六项复核失败、无空 FOUP、MES 分批失败、Transfer FOUP 失败。",
        "Litho、BARCO、其他站点的 WaitPilotChangeFOUP 解除条件。",
        "子母批同机台与 Specify Lot 回归。",
    ])

    approval_heading = add_heading(doc, "十一、审批意见", 1)
    approval_heading.paragraph_format.page_break_before = True
    add_approval_table(doc)

    doc.core_properties.title = "LithoAutoPiRun 需求申请单（最终版）"
    doc.core_properties.subject = "物理分批、Pilot 选择、Transfer FOUP 及相关卡控逻辑"
    doc.core_properties.author = "制造部"
    doc.core_properties.keywords = "LithoAutoPiRun, RTD, AMA, Pilot, FOUP, 物理分批"

    force_all_text_black(doc)
    apply_body_typography(doc)
    doc.save(OUT_PATH)
    shutil.copy2(OUT_PATH, QA_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
