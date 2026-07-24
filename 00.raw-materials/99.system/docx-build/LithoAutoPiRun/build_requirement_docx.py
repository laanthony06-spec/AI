from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Obsidian\work\OBSidianCodex")
OUT_DIR = ROOT / "00.raw-materials" / "90.processed" / "LithoAutoPiRun"
BUILD_DIR = ROOT / "00.raw-materials" / "99.system" / "docx-build" / "LithoAutoPiRun"
ASSET_DIR = BUILD_DIR / "assets"
OUT_PATH = OUT_DIR / "LithoAutoPiRun_新增需求申请单_修订稿.docx"

OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "AEB8C4"
WHITE = "FFFFFF"
BLACK = "111827"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_run_font(run, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_run(run, fill="EEF2F6"):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def add_rich_text(paragraph, text: str, size=11, color=BLACK, bold=False):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, ascii_font="Consolas", east_asia="Microsoft YaHei", size=size - 0.3,
                         color=DARK_BLUE, bold=False)
            shade_run(run)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_field(paragraph, instruction: str, display="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def create_numbering(doc: Document, kind="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def add_body(doc, text, bold=False, align=None, after=6, keep=False):
    p = doc.add_paragraph(style="Normal")
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if keep:
        p.paragraph_format.keep_with_next = True
    add_rich_text(p, text, bold=bold)
    return p


def add_bullets(doc, items):
    num_id = create_numbering(doc, "bullet")
    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        add_rich_text(p, item)


def add_numbers(doc, items):
    num_id = create_numbering(doc, "decimal")
    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        add_rich_text(p, item)


def add_code_block(doc, text):
    p = doc.add_paragraph(style="Code Block")
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, ascii_font="Consolas", east_asia="Microsoft YaHei", size=9.2, color=INK_BLUE)
    return p


def set_style_font(style, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None,
                   color=None, bold=None):
    style.font.name = ascii_font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def setup_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, INK_BLUE, 6, 3),
    ):
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Code Block" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Code Block"]
    set_style_font(style, ascii_font="Consolas", east_asia="Microsoft YaHei", size=9.2, color=INK_BLUE)
    style.paragraph_format.left_indent = Inches(0.18)
    style.paragraph_format.right_indent = Inches(0.10)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.08
    ppr = style._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    ppr.append(shd)

    if "Figure Caption" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Figure Caption"]
    set_style_font(style, size=9.5, color=MUTED, bold=True)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.keep_with_next = True


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    return section


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = p.add_run("LithoAutoPiRun 功能需求规格")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = p.add_run("\t修订稿 | 2026-07-24")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(p, "PAGE", "1")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        set_run_font(r, size={1:16, 2:13, 3:12, 4:11}[level],
                     color={1:BLUE, 2:BLUE, 3:DARK_BLUE, 4:INK_BLUE}[level], bold=True)
    return p


def add_metadata_table(doc):
    rows = [
        ("编号", "由信息技术部填写", "类别", "功能开发"),
        ("申请部门", "制造部", "申请人员", "温浩奇"),
        ("系统名称", "CIM 计算机集成制造系统 Fab6（一科）", "功能模块", "智能派工系统（RTD/DSP）"),
        ("申请日期", "2026-07-24", "希望交付日期", "2026-07-29"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1400, 3280, 1400, 3280])
    set_table_borders(table)
    for row, values in zip(table.rows, rows):
        set_cant_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, value, size=9.7, bold=(idx % 2 == 0))
            if idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_approval_table(doc):
    rows = [
        "申请部门意见",
        "申请部门分管领导意见",
        "相关部门意见",
        "相关部门分管领导意见",
        "信息技术部意见",
        "信息技术部分管领导意见",
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    for row, label in zip(table.rows, rows):
        set_cant_split(row)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_rich_text(p, label, size=10, bold=True)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(34)
        add_rich_text(p2, "意见：\n\n日期：", size=10)


def find_font(size):
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for para in text.split("\n"):
        if not para:
            lines.append("")
            continue
        current = ""
        for char in para:
            candidate = current + char
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill, max_width_pad=28):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, (x2 - x1) - max_width_pad * 2)
    bbox = draw.textbbox((0, 0), "中A", font=font)
    line_h = bbox[3] - bbox[1] + 9
    total = line_h * len(lines)
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        tb = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (tb[2] - tb[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def draw_node(draw, node, font):
    x, y, w, h, text, kind, fill = node
    box = (x, y, x+w, y+h)
    outline = (65, 92, 120)
    if kind == "diamond":
        points = [(x+w/2, y), (x+w, y+h/2), (x+w/2, y+h), (x, y+h/2)]
        draw.polygon(points, fill=fill, outline=outline)
        draw.line(points + [points[0]], fill=outline, width=3)
        inner = (x+w*0.15, y+h*0.18, x+w*0.85, y+h*0.82)
        draw_centered_text(draw, inner, text, font, (20, 35, 55), 10)
    elif kind == "end":
        draw.rounded_rectangle(box, radius=h//2, fill=fill, outline=outline, width=3)
        draw_centered_text(draw, box, text, font, (20, 35, 55))
    else:
        draw.rounded_rectangle(box, radius=20, fill=fill, outline=outline, width=3)
        draw_centered_text(draw, box, text, font, (20, 35, 55))


def node_anchor(node, side):
    x, y, w, h, *_ = node
    return {
        "top": (x+w/2, y), "bottom": (x+w/2, y+h),
        "left": (x, y+h/2), "right": (x+w, y+h/2),
    }[side]


def draw_arrow(draw, start, end, label=None, font=None, via=None):
    pts = [start] + (via or []) + [end]
    draw.line(pts, fill=(70, 85, 105), width=4, joint="curve")
    x2, y2 = pts[-1]
    x1, y1 = pts[-2]
    angle = math.atan2(y2-y1, x2-x1)
    size = 13
    p1 = (x2 - size*math.cos(angle-0.5), y2 - size*math.sin(angle-0.5))
    p2 = (x2 - size*math.cos(angle+0.5), y2 - size*math.sin(angle+0.5))
    draw.polygon([(x2, y2), p1, p2], fill=(70, 85, 105))
    if label and font:
        mx, my = pts[len(pts)//2]
        tb = draw.textbbox((0, 0), label, font=font)
        draw.rounded_rectangle((mx-8, my-8, mx+(tb[2]-tb[0])+10, my+(tb[3]-tb[1])+8),
                               radius=6, fill=(255, 255, 255))
        draw.text((mx, my), label, font=font, fill=(40, 55, 75))


def make_flowcharts():
    font = find_font(27)
    small = find_font(23)
    bg = (250, 252, 255)
    process = (232, 238, 245)
    decision = (255, 247, 224)
    success = (229, 243, 235)
    stop = (253, 232, 232)

    # 1. RTD
    path = ASSET_DIR / "01-rtd-selection.png"
    img = Image.new("RGB", (1800, 2200), bg)
    d = ImageDraw.Draw(img)
    nodes = {
        "a": (330, 45, 930, 115, "从 FAB6、FAB8 获取候选 Lot", "end", process),
        "b": (330, 210, 930, 125, "获取 Lot、Carrier、Capability、RemainQ 等基础信息", "process", process),
        "c": (420, 390, 750, 155, "满足基础筛选条件？", "diamond", decision),
        "d": (330, 600, 930, 125, "向后 Fetch 20 个站点并形成 PiRunLoop", "process", process),
        "e": (420, 780, 750, 155, "存在有效 CD、Litho 和 Reticle？", "diamond", decision),
        "f": (420, 990, 750, 170, "同 FOUP 已有 Pilot、FutureHold 或 RC？", "diamond", decision),
        "g": (330, 1215, 930, 125, "获取并筛选可作业 Litho 机台", "process", process),
        "h": (420, 1395, 750, 160, "存在 Pi_split_flag='Y' 的有效机台？", "diamond", decision),
        "i": (330, 1610, 930, 125, "检查 R2R Context，剔除多路径 Lot", "process", process),
        "j": (330, 1790, 930, 125, "计算 Lot 与 Context 排序指标", "process", process),
        "k": (330, 1970, 930, 125, "循环选择每个 Context 的最优 Pilot", "end", success),
        "x": (1370, 920, 320, 130, "剔除 Lot", "end", stop),
    }
    for n in nodes.values(): draw_node(d, n, font)
    for s, t in (("a","b"),("b","c"),("d","e"),("g","h"),("i","j"),("j","k")):
        draw_arrow(d, node_anchor(nodes[s], "bottom"), node_anchor(nodes[t], "top"))
    draw_arrow(d, node_anchor(nodes["c"], "bottom"), node_anchor(nodes["d"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["c"], "right"), node_anchor(nodes["x"], "top"), "否", small,
               [(1300, 468), (1530, 468)])
    draw_arrow(d, node_anchor(nodes["e"], "bottom"), node_anchor(nodes["f"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["e"], "right"), node_anchor(nodes["x"], "left"), "否", small,
               [(1270, 858), (1270, 985)])
    draw_arrow(d, node_anchor(nodes["f"], "bottom"), node_anchor(nodes["g"], "top"), "否", small)
    draw_arrow(d, node_anchor(nodes["f"], "right"), node_anchor(nodes["x"], "left"), "是", small)
    draw_arrow(d, node_anchor(nodes["h"], "bottom"), node_anchor(nodes["i"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["h"], "right"), node_anchor(nodes["x"], "bottom"), "否", small,
               [(1310, 1475), (1530, 1475)])
    img.save(path, quality=95)

    # 2. Pilot decision
    path = ASSET_DIR / "02-pilot-split.png"
    img = Image.new("RGB", (1800, 2050), bg)
    d = ImageDraw.Draw(img)
    nodes = {
        "a": (380, 45, 900, 115, "取得已选 Lot + Context", "end", process),
        "b": (450, 225, 760, 170, "满足任一整批 Pilot 条件？", "diamond", decision),
        "c": (70, 480, 520, 130, "IsNeedSplit=F\n整批 Lot 设置为 Pilot", "end", success),
        "d": (760, 480, 760, 120, "IsNeedSplit=T，读取 pi_splitcnt", "process", process),
        "e": (760, 665, 760, 155, "pi_splitcnt 为空、≤0 或 >25？", "diamond", decision),
        "f": (1210, 885, 470, 115, "使用默认值 4", "process", process),
        "g": (620, 1050, 760, 160, "选片数大于当前可用 Wafer 数？", "diamond", decision),
        "h": (620, 1270, 760, 125, "按 Wafer ID、Chuck/Slot 建立分组", "process", process),
        "i": (620, 1455, 760, 125, "计算 GroupRank 和 WaferRank", "process", process),
        "j": (620, 1640, 760, 135, "按 WaferRank ASC、GroupRank ASC 选片", "process", process),
        "k": (620, 1835, 760, 125, "生成 pi_splitwafer 并确定 Merge 站点", "end", success),
    }
    for n in nodes.values(): draw_node(d, n, font)
    draw_arrow(d, node_anchor(nodes["a"], "bottom"), node_anchor(nodes["b"], "top"))
    draw_arrow(d, node_anchor(nodes["b"], "left"), node_anchor(nodes["c"], "top"), "是", small,
               [(330, 310), (330, 445)])
    draw_arrow(d, node_anchor(nodes["b"], "bottom"), node_anchor(nodes["d"], "top"), "否", small)
    draw_arrow(d, node_anchor(nodes["d"], "bottom"), node_anchor(nodes["e"], "top"))
    draw_arrow(d, node_anchor(nodes["e"], "bottom"), node_anchor(nodes["g"], "top"), "否，使用配置值", small)
    draw_arrow(d, node_anchor(nodes["e"], "right"), node_anchor(nodes["f"], "top"), "是", small,
               [(1590, 742)])
    draw_arrow(d, node_anchor(nodes["f"], "bottom"), node_anchor(nodes["g"], "right"))
    draw_arrow(d, node_anchor(nodes["g"], "left"), node_anchor(nodes["c"], "bottom"), "是，改为整批", small,
               [(330, 1130)])
    draw_arrow(d, node_anchor(nodes["g"], "bottom"), node_anchor(nodes["h"], "top"), "否", small)
    for s, t in (("h","i"),("i","j"),("j","k")):
        draw_arrow(d, node_anchor(nodes[s], "bottom"), node_anchor(nodes[t], "top"))
    img.save(path, quality=95)

    # 3. WaitPilot control
    path = ASSET_DIR / "03-wait-pilot-control.png"
    img = Image.new("RGB", (1800, 1900), bg)
    d = ImageDraw.Draw(img)
    nodes = {
        "a": (380, 45, 900, 115, "获取需要 Transfer FOUP 的 Litho Pilot", "end", process),
        "b": (440, 220, 780, 165, "WatchDog 已开启且位于触发时间范围？", "diamond", decision),
        "z": (1360, 245, 330, 115, "不新增卡控", "end", success),
        "c": (440, 465, 780, 165, "Pilot 位于 UnscheduledSorter？", "diamond", decision),
        "d": (100, 720, 660, 130, "卡控同 FOUP 中不在 Sorter 的 Other Lot", "process", process),
        "e": (1040, 720, 660, 130, "卡控同 FOUP 中同样不在 Sorter 的 Other Lot", "process", process),
        "f": (440, 940, 780, 165, "Other Lot 当前站点？", "diamond", decision),
        "g": (60, 1210, 440, 120, "Litho：不解除卡控", "end", stop),
        "h": (680, 1190, 520, 155, "BARCO：RemainQ<4 h\n或触发 Qu_0？", "diamond", decision),
        "i": (1320, 1190, 420, 155, "其他站点：\n触发 Qu_0？", "diamond", decision),
        "j": (680, 1510, 520, 120, "解除 WaitPilotChangeFOUP", "end", success),
        "k": (680, 1710, 520, 110, "否则保持卡控", "end", stop),
    }
    for n in nodes.values(): draw_node(d, n, font)
    draw_arrow(d, node_anchor(nodes["a"], "bottom"), node_anchor(nodes["b"], "top"))
    draw_arrow(d, node_anchor(nodes["b"], "right"), node_anchor(nodes["z"], "left"), "否", small)
    draw_arrow(d, node_anchor(nodes["b"], "bottom"), node_anchor(nodes["c"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["c"], "left"), node_anchor(nodes["d"], "top"), "是", small,
               [(430, 548), (430, 680)])
    draw_arrow(d, node_anchor(nodes["c"], "right"), node_anchor(nodes["e"], "top"), "否", small,
               [(1370, 548), (1370, 680)])
    draw_arrow(d, node_anchor(nodes["d"], "bottom"), node_anchor(nodes["f"], "left"), via=[(430, 1022)])
    draw_arrow(d, node_anchor(nodes["e"], "bottom"), node_anchor(nodes["f"], "right"), via=[(1370, 1022)])
    draw_arrow(d, node_anchor(nodes["f"], "left"), node_anchor(nodes["g"], "top"), "Litho", small,
               [(280, 1022), (280, 1170)])
    draw_arrow(d, node_anchor(nodes["f"], "bottom"), node_anchor(nodes["h"], "top"), "BARCO", small)
    draw_arrow(d, node_anchor(nodes["f"], "right"), node_anchor(nodes["i"], "top"), "其他", small,
               [(1530, 1022), (1530, 1150)])
    draw_arrow(d, node_anchor(nodes["h"], "bottom"), node_anchor(nodes["j"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["i"], "bottom"), node_anchor(nodes["j"], "right"), "是", small,
               [(1530, 1570)])
    draw_arrow(d, node_anchor(nodes["h"], "left"), node_anchor(nodes["k"], "left"), "否", small,
               [(580, 1268), (580, 1765)])
    draw_arrow(d, node_anchor(nodes["i"], "left"), node_anchor(nodes["k"], "right"), "否", small,
               [(1260, 1268), (1260, 1765)])
    img.save(path, quality=95)

    # 4. AMA
    path = ASSET_DIR / "04-ama-split.png"
    img = Image.new("RGB", (1800, 2050), bg)
    d = ImageDraw.Draw(img)
    nodes = {
        "a": (390, 45, 900, 115, "AMA 读取 Central_GetLithoR2RAutoPirunInfo", "end", process),
        "b": (470, 220, 740, 155, "IsNeedSplit=T？", "diamond", decision),
        "c": (60, 500, 500, 120, "整批 Lot 传给 R2R", "end", success),
        "d": (760, 485, 760, 130, "执行物理分批前六项复核", "process", process),
        "e": (760, 680, 760, 155, "六项复核全部通过？", "diamond", decision),
        "f": (70, 900, 520, 135, "停止处理并记录原因\n等待下一轮重新计算", "end", stop),
        "g": (760, 900, 760, 125, "获取并预占空 FOUP", "process", process),
        "h": (760, 1085, 760, 155, "是否成功取得空 FOUP？", "diamond", decision),
        "i": (760, 1300, 760, 120, "调用 MES 物理分批接口", "process", process),
        "j": (760, 1480, 760, 155, "接口调用成功？", "diamond", decision),
        "k": (80, 1720, 520, 125, "释放预占 FOUP\n整批 Lot 传给 R2R", "end", stop),
        "l": (1040, 1710, 660, 140, "生成子批并设置为 Pilot\n按现有顺序传给 R2R、执行 Transfer FOUP", "end", success),
    }
    for n in nodes.values(): draw_node(d, n, font)
    draw_arrow(d, node_anchor(nodes["a"], "bottom"), node_anchor(nodes["b"], "top"))
    draw_arrow(d, node_anchor(nodes["b"], "left"), node_anchor(nodes["c"], "top"), "否", small,
               [(310, 298), (310, 460)])
    draw_arrow(d, node_anchor(nodes["b"], "bottom"), node_anchor(nodes["d"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["d"], "bottom"), node_anchor(nodes["e"], "top"))
    draw_arrow(d, node_anchor(nodes["e"], "left"), node_anchor(nodes["f"], "top"), "否", small,
               [(330, 758), (330, 860)])
    draw_arrow(d, node_anchor(nodes["e"], "bottom"), node_anchor(nodes["g"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["g"], "bottom"), node_anchor(nodes["h"], "top"))
    draw_arrow(d, node_anchor(nodes["h"], "left"), node_anchor(nodes["c"], "bottom"), "否，整批回退", small,
               [(310, 1162), (310, 660)])
    draw_arrow(d, node_anchor(nodes["h"], "bottom"), node_anchor(nodes["i"], "top"), "是", small)
    draw_arrow(d, node_anchor(nodes["i"], "bottom"), node_anchor(nodes["j"], "top"))
    draw_arrow(d, node_anchor(nodes["j"], "left"), node_anchor(nodes["k"], "top"), "否", small,
               [(340, 1558), (340, 1680)])
    draw_arrow(d, node_anchor(nodes["j"], "right"), node_anchor(nodes["l"], "top"), "是", small,
               [(1370, 1645)])
    img.save(path, quality=95)


def add_figure(doc, caption, image_path, alt_text, width_inches=6.15):
    p = doc.add_paragraph(caption, style="Figure Caption")
    p.paragraph_format.keep_with_next = True
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(8)
    pic_p.paragraph_format.keep_together = True
    run = pic_p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width_inches))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)


def build_document():
    make_flowcharts()
    doc = Document()
    setup_styles(doc)
    section = setup_page(doc)
    setup_header_footer(section)

    # memo_masthead opening
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("功能需求规格")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("LithoAutoPiRun 新增需求申请单")
    set_run_font(r, size=24, color=INK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("逻辑分批改造为物理分批及 Pilot 选择、Transfer FOUP、卡控逻辑优化")
    set_run_font(r, size=12.5, color=MUTED)

    add_metadata_table(doc)

    add_heading(doc, "一、项目简介及必要性分析", 1)
    add_body(doc, "当前 `LithoAutoSplitPirun` 采用逻辑分批方式，生成的 Pilot 与母批仍位于同一 FOUP。在跨厂转运场景下，由于同一 FOUP 内 Lot 的转运限制，Pilot 可能无法及时执行 PiRun，进而影响产线 WIP 流转。")
    add_body(doc, "现有 AutoPiRun 仅会为满足分批条件的 Lot 自动设置 Pilot。不满足条件的 Lot 可能持续处于卡控状态，增加 Queue Time 超时风险。")
    add_body(doc, "本需求拟将 Pilot 的拆分方式由逻辑分批调整为物理分批，同时优化候选 Lot 筛选、Pilot 排序、Transfer FOUP 以及子母批同机台卡控逻辑。")

    add_heading(doc, "二、项目投资方案比较及效果分析", 1)
    add_heading(doc, "改善方案", 2)
    add_body(doc, "将 `LithoAutoSplitPirun` 由逻辑分批改为物理分批，并优化 Pilot 的筛选、排序、分批、Transfer FOUP 及卡控逻辑。")
    add_heading(doc, "预期效果", 2)
    add_body(doc, "降低 Lot Queue Time 超时风险。", bold=True)

    add_heading(doc, "三、需求内容", 1)
    add_heading(doc, "一、RTD 新增逻辑", 2)
    add_body(doc, "本次 RTD 调整包括：")
    add_numbers(doc, [
        "修改 Report `Central_GetLithoR2RAutoPirunInfo` 中的 Pilot 选择逻辑。",
        "新增 Report `LithoPilotAutoDoAdhocSorter`。",
        "在 Global Macro、LithoRule 和 LithoAssign 中增加相关卡控逻辑。",
        "`Pi_split_flag` 统一使用 `Y`，表示允许物理分批。",
        "PiRun 状态统一使用 `PIRUNON`。",
    ])

    add_heading(doc, "1. 修改 Report：Central_GetLithoR2RAutoPirunInfo", 2)
    add_heading(doc, "1.1 Lot 获取", 3)
    add_body(doc, "从 FAB6 和 FAB8 获取待判断的 Lot，并按照以下规则逐步筛选。")
    add_heading(doc, "1.1.1 Lot 基础信息获取", 4)
    add_bullets(doc, [
        "从 `fwlot` 获取 `appid`、`priority`、`processingstatus`、`componentqty`。",
        "从 `fablotext` 获取 `requiredcapability`、`runcardid`、`reticleid`。",
        "从 `fablotcarrierext` 获取 `carrierkind`。",
        "从 `fabinqtimeprocess` 获取 `RemainQ`。",
        "从 `UI.RTDConfig-LITHOLotAssignment-LithoAssignCapability` 获取 `LithoCapability`。",
    ])
    add_body(doc, "BARCO Capability 固定为：", keep=True)
    add_code_block(doc, "L-BARCO-L\nL-BARCO-S")

    add_heading(doc, "1.1.2 Lot 基础筛选条件", 4)
    add_body(doc, "保留同时满足以下基础条件的 Lot：", keep=True)
    add_code_block(doc, """processingstatus IN ('Active', 'CrossFabTransferred')
AND carrierkind = 'FOUP'
AND runcardid IS NULL
AND requiredcapability IN (
    LithoCapability,
    'L-BARCO-L',
    'L-BARCO-S'
)""")
    add_body(doc, "跨厂数据继续使用以下现有筛选表达式：", keep=True)
    add_code_block(doc, """IsTransferLot = True
OR (
    IsTransferLot <> True
    AND processingstatus <> 'CrossFabTransferred'
)""")
    add_body(doc, "`IsTransferLot` 由现有 `istransferlot` Macro 判断。本需求不增加额外的 FAB6/FAB8 跨厂数据去重规则。")

    add_heading(doc, "1.2 获取 PiRun 站点并进一步筛选 Lot", 3)
    add_heading(doc, "1.2.1 获取 PiRunLoop", 4)
    add_body(doc, "对通过 1.1 筛选的 Lot，向后 Fetch 20 个站点，并获取 `productname`、`planname`、`stage`、`capability`、`stepseq`、`recipeid`、`STN`。Transfer Lot 按现有逻辑修正厂别信息。")
    add_body(doc, "将获取的站点截取到相同 Stage 的最后一道 CD 站点。Lot 当前站点至最后一道 CD 站点之间的区段定义为 `PiRunLoop`。如果 Fetch 结果中不存在 CD 站点，则剔除该 Lot。现有流程负责保证后续能够获得有效 Merge 站点，本需求不增加额外异常处理。")

    add_heading(doc, "1.2.2 其他筛选条件", 4)
    add_numbers(doc, [
        "Litho 站点判断：`PiRunLoop` 中必须存在 `capability=LithoCapability` 且具有 Reticle 信息的站点；不满足时剔除该 Lot。",
        "Specify Lot 判断：根据 `productid`、`layer`、`lotid` 匹配 `r2r_litho_whitelist`。匹配成功时按现有 Specify Lot 规则处理。",
        "同 FOUP Pilot 判断：从 `rtd_r2r_litho_context_ovlcd` 和 `rtd_r2r_lot_history` 获取 Litho Pilot 信息。同一 FOUP 中已经存在 Litho Pilot 时，剔除当前候选 Lot。",
        "FutureHold 判断：从 `fabfutureaction` 获取 `PiRunLoop` 中各站点的 FutureAction。存在 FutureHold 时，剔除该 Lot。",
        "RC 判断：`PiRunLoop` 中存在 RC 站点时，剔除该 Lot。",
    ])

    add_heading(doc, "1.3 R2R 条件判断", 3)
    add_body(doc, "获取 Lot 在可作业机台上的 R2R 状态，并判断该 Lot 是否存在多路径。")
    add_heading(doc, "1.3.1 可作业机台获取", 4)
    add_body(doc, "获取 1.2 中 Lot 的 Litho 站点。通过现有 Transfer Macro 判断 Lot 是否能够在对应 Litho 站点执行 Transfer；若可以，则同时获取对应工厂的 Litho 机台。")
    add_body(doc, "按厂别从 `rtd_r2r_litho_add_setting` 获取机台的 `Pi_split_flag` 和 `pi_splitcnt`。Lot 至少需要存在一台 `Pi_split_flag='Y'` 的可作业机台，否则剔除该 Lot。")
    add_body(doc, "对剩余机台继续检查 EQPStatus、LCC、Capability、Recipe、PPID 和 Global Reason。存在卡控时，剔除对应机台。")

    add_heading(doc, "1.3.2 多路径判断", 4)
    add_body(doc, "按照 `Lot + STN + Reticle` 维度，从 `rtd_r2r_litho_context_ovl` 和 `rtd_r2r_litho_context_cd` 匹配 R2R Context。保留满足以下条件且不存在 R2R Reason 的 Context：", keep=True)
    add_code_block(doc, """OVL_Status IN ('PIRUNON', 'ON', 'Fixed')
AND CD_Status IN ('PIRUNON', 'ON', 'Fixed')""")
    add_body(doc, "按 Lot 统计有效 Context 数量，记为 `ContextCount`。当 `ContextCount>1` 时，判定该 Lot 存在多路径，并剔除该 Lot。")

    add_heading(doc, "1.4 选择 Lot", 3)
    add_body(doc, "按 Context 对候选 Lot 排序，并循环选择最优的 `Lot + Context`。")
    add_heading(doc, "1.4.1 AutoPiRun Context 筛选", 4)
    add_body(doc, "获取 1.3 筛选后的 Lot 及 Context 信息，包括 Lot、STN、Reticle、Prod、Layer、Recipe、Pretool、Prereticle、`Custom_Context_Value`、`CD_Status`、`OVL_Status`、`Pilot_CD`、`Pilot_OVL`、`Pi_split_flag` 和 `pi_splitcnt`。")
    add_body(doc, "保留满足以下条件的 Context：", keep=True)
    add_code_block(doc, """Pi_split_flag = 'Y'
AND (
    Pilot_CD IS NULL
    OR Pilot_OVL IS NULL
)""")

    add_heading(doc, "1.4.2 Context 内 Lot 排序", 4)
    add_numbers(doc, [
        "GapToLitho：计算 Lot 当前站点距离 Litho 站点的剩余 Step 数量。",
        "SplitCntMatched：`componentqty>=pi_splitcnt` 时取 1，否则取 0。`pi_splitcnt` 为空、等于 0、小于 0 或大于 25 时，使用默认值 4；配置值在 1～25 范围内但大于 Lot 当前可用 Wafer 数时，后续按整批 Pilot 处理。",
        "RequiredChuckCount：Lot 存在 `prelayer` 时，从 `r2r_litho_waferhistory` 获取各 Wafer 的 Chuck 信息；否则从 `fsmaterialassociation` 获取 Wafer Slot 信息。C1、C2 各至少包含 2 片，或奇偶 Slot 各至少包含 2 片时取 1，否则取 0。Chuck 信息由上游保证有效和完整。",
        "BulletLot：按现有逻辑判断 Lot 是否为空扣或空 LP Lot；是取 1，否则取 0。",
        "RemainQ：使用实际剩余 Queue Time；空值赋为 `9999 h`。",
        "KeyLot：`quota_applyinfo` 中 `KeyLot=1 AND Status='CONFIRM'` 时取 1，否则取 0。",
    ])
    add_body(doc, "按以下优先级对 Lot 排序，结果记为 `RTDRank`：", keep=True)
    add_code_block(doc, """1. GapToLitho ASC
2. SplitCntMatched DESC
3. RequiredChuckCount DESC
4. BulletLot DESC
5. RemainQ ASC
6. KeyLot ASC
7. lotid ASC""")
    add_body(doc, "其中 `KeyLot ASC` 表示普通 Lot 优先于 Key Lot。")

    add_heading(doc, "1.4.3 Context 之间排序", 4)
    add_numbers(doc, [
        "ReticleSTNRank：按 `Reticle + STN` 对 Context 分组，并获取 Reticle 当前所在机台 `ReticleOnSTN`。第一轮中，`STN=ReticleOnSTN` 时取 1，否则取 0；后续循环中，与上一轮选中 Context 属于同一 `Reticle + STN` 分组时取 1，否则取 0。",
        "ContextCandidateCount：统计当前 Context 中的候选 Lot 数量。",
        "ActualSTNPilotCount：按 STN 统计已经选中的 Pilot Context 数量。",
    ])
    add_body(doc, "最终按照以下优先级对 `Lot + Context` 排序：", keep=True)
    add_code_block(doc, """1. ReticleSTNRank DESC
2. ContextCandidateCount ASC
3. ActualSTNPilotCount ASC
4. RTDRank ASC""")

    add_heading(doc, "1.4.4 循环选择 Pilot", 4)
    add_body(doc, "将排序第一的 `Lot + Context` 固定为已选 Pilot Context。每个 Context 最多选择一个 Pilot；选中后，剔除所有与该已选 Context 相同的其他候选 `Lot + Context`。")
    add_body(doc, "每选中一个 Pilot 后，立即重新计算 `ReticleSTNRank`、`ContextCandidateCount` 和 `ActualSTNPilotCount`，完成指标更新后进入下一轮循环，直至没有可用 Context 或没有可用 Lot。")

    add_figure(doc, "图 1  RTD 候选筛选与 Pilot 选择流程", ASSET_DIR / "01-rtd-selection.png",
               "RTD 从 FAB6、FAB8 获取候选 Lot，依次完成基础筛选、站点检查、机台检查、R2R Context 检查和循环选择 Pilot。")

    add_heading(doc, "1.5 Pilot 选片逻辑", 3)
    add_body(doc, "判断 Lot 是否需要物理分批。不满足物理分批条件时，将整批设置为 Pilot。")
    add_heading(doc, "1.5.1 整批设置为 Pilot", 4)
    add_body(doc, "以下五项条件使用“或”关系。任意一项满足时，将整批 Lot 设置为 Pilot：", keep=True)
    add_code_block(doc, """1. BulletLot=1 OR KeyLot=1
OR
2. CurCapability=LithoCapability AND FuLL(RemainQ)
OR
3. RequiredChuckCount=0
OR
4. SplitCntMatched=0
OR
5. componentqty<=6""")
    add_body(doc, "`componentqty>0` 由上游系统保证，本需求不增加零片判断。满足任一条件时，设置 `IsNeedSplit=F`。")

    add_heading(doc, "1.5.2 物理分批选片逻辑", 4)
    add_body(doc, "不满足 1.5.1 任一整批条件时，设置 `IsNeedSplit=T`，并执行物理分批选片。")
    add_heading(doc, "1.5.2.1 Wafer 分组", 4)
    add_code_block(doc, """Wafer ID 1～10  → Group1
Wafer ID 11～25 → Group2

存在 Chuck 信息：C1→SubGroup1，C2→SubGroup2
不存在 Chuck 信息：奇数 Slot→SubGroup1，偶数 Slot→SubGroup2

Group1 + SubGroup1 → GroupRank=1
Group1 + SubGroup2 → GroupRank=2
Group2 + SubGroup1 → GroupRank=3
Group2 + SubGroup2 → GroupRank=4""")

    add_heading(doc, "1.5.2.2 Wafer 排序与选择", 4)
    add_body(doc, "在每个 `Group + SubGroup` 内，按 Wafer ID 从小到大排序并编号，记为 `WaferRank`。最终按 `WaferRank ASC`、`GroupRank ASC` 选择 Wafer，用于从各 Group/SubGroup 轮流、均匀选片。")
    add_bullets(doc, [
        "`pi_splitcnt` 为空、等于 0、小于 0 或大于 25 时，使用默认值 4。",
        "`pi_splitcnt` 在 1～25 范围内时，按配置数量选择。",
        "`pi_splitcnt` 大于 Lot 当前可用 Wafer 数时，不再物理分批，改为整批 Pilot。",
        "选中的 Wafer 记为 `pi_splitwafer`；物理分批后，将 `pi_splitwafer` 对应的子批设置为 Pilot。",
    ])

    add_heading(doc, "1.5.3 Merge 站点设置", 4)
    add_body(doc, "如果 Lot 存在 ADI 站点，且该 ADI 站点不包含 SRC，则将第一道符合条件的 ADI 设置为 Merge 站点；否则，将最后一道 CD 设置为 Merge 站点。有效 Merge 站点由现有流程保证，本需求不增加两者均不存在时的异常处理。")

    add_figure(doc, "图 2  Pilot 整批与物理分批判断流程", ASSET_DIR / "02-pilot-split.png",
               "Pilot 先判断任一整批条件，再校验 pi_splitcnt，并按照 Group、SubGroup、WaferRank 和 GroupRank 选择物理分批 Wafer。")

    add_heading(doc, "1.6 输出 Report", 3)
    add_body(doc, "将现有计算结果存入 `Central_GetLithoR2RAutoPirunInfo`。Report 字段包括：")
    add_bullets(doc, [
        "Lot、`toolid`、`productid`、`layerid`、`reticleid`",
        "`prereticle`、`pretool`、`custom_context_value`",
        "`pi_splitwafer`、`IsNeedSplit`、`isSTNSite`",
    ])
    add_body(doc, "保持现有 Report 输出范围和记录方式，不新增 `request_id`、`generated_time`、`execution_status` 或 `error_code`。")

    add_heading(doc, "2. 新增 Report：LithoPilotAutoDoAdhocSorter", 2)
    add_body(doc, "该 Report 用于获取需要 Transfer FOUP 的 Litho Pilot。")
    add_heading(doc, "2.1 Litho Pilot 获取", 3)
    add_heading(doc, "2.1.1 数据获取", 4)
    add_bullets(doc, [
        "从 `r2r_litho_context_ovl` 获取 `ovl_status`、`pilot`。",
        "从 `r2r_litho_context_cd` 获取 `cd_status`、`pilot`。",
        "从 `fwlot` 获取 `appid`、`lottype`、`priority`。",
        "从 `fabcategorymap` 获取 `lottype`、`category`。",
    ])
    add_heading(doc, "2.1.2 筛选条件", 4)
    add_code_block(doc, """priority < 5
AND category = 'Production'
AND (
    ovl_status = 'PIRUNON'
    OR cd_status = 'PIRUNON'
)""")

    add_heading(doc, "2.2 Transfer FOUP 判断", 3)
    add_body(doc, "获取 2.1 中 Litho Pilot 的 Carrier 信息。按 Carrier 从 `fwlot` 获取所有 Lot，并保留 `fwlot.extrastatus='WaitForJobPrep'` 的记录。如果 Litho Pilot 所在 Carrier 中存在其他 Lot，则该 Pilot 需要 Change FOUP。")

    add_heading(doc, "2.3 Carrier 排序规则", 3)
    add_body(doc, "对需要 Transfer FOUP 的 Litho Pilot，按以下顺序排序，并根据排序结果建立 `AdhocSorterJob`：", keep=True)
    add_code_block(doc, """1. RemainQ ASC
2. Priority ASC
3. componentqty DESC
4. lotid ASC""")

    add_heading(doc, "2.4 输出 Report", 3)
    add_body(doc, "从 `AMA.TriggerConfig-WatchDog_LithoPilotAutoDoAdhocSorter` 获取 `Switch`、`Trigger Time Slot` 和 `TriggerCount/Time`。当 `Switch='Y'` 且当前时间位于现有 WatchDog 的 Trigger Time Slot 判断范围内时，按照 `TriggerCount/Time` 选择需要物理分批的 Carrier，并将结果存入 `LithoPilotAutoDoAdhocSorter`。")
    add_body(doc, "`Trigger Time Slot` 沿用 WatchDog 现有判断方式，本需求不增加时区或区间边界定义。Report 字段包括 Carrier、Pilot、`extrastatus`、Status、RemainQ、Pieces、Prod、Priority。")

    add_heading(doc, "3. Rule 中新增卡控逻辑", 2)
    add_heading(doc, "3.1 Global Macro 新增卡控逻辑", 3)
    add_body(doc, "在 Global Macro 中增加对需要 Transfer FOUP 的 Litho Pilot 的卡控。当 `WatchDog_LithoPilotAutoDoAdhocSorter.Switch='Y'` 且当前时间位于现有 Trigger Time Slot 范围内时，按照以下场景处理。")
    add_heading(doc, "3.1.1 Pilot 位于 AdhocSorter 站点", 4)
    add_body(doc, "如果 `adhocplanname` 包含 `UnscheduledSorter`，则同一 FOUP 中不在 AdhocSorter 站点的 Other Lot 增加 `Reason=WaitPilotChangeFOUP`。解除规则如下：")
    add_numbers(doc, [
        "Other Lot 位于 Litho 站点时，不解除卡控。",
        "Other Lot 位于 BARCO 站点时，`RemainQ<4 h` 或触发 `Qu_0`，任一条件满足即可解除。",
        "Other Lot 位于非 Litho、非 BARCO 站点时，仅在触发 `Qu_0` 时解除。",
    ])

    add_heading(doc, "3.1.2 Pilot 不在 AdhocSorter 站点", 4)
    add_body(doc, "如果 Litho Pilot 不在 AdhocSorter 站点，则同一 FOUP 中同样不在 AdhocSorter 站点的 Other Lot 增加 `Reason=WaitPilotChangeFOUP`。Other Lot 的解除规则沿用 3.1.1。")
    add_body(doc, "当 Litho Pilot 的 `RemainQ<4 h` 或 Pilot 触发 `Qu_0` 时，解除相关卡控。Other Lot 紧急解除卡控后，Pilot 的 Transfer FOUP 任务由现有 Adhoc Sorter 流程处理，本需求不增加额外规则。")

    add_figure(doc, "图 3  WaitPilotChangeFOUP 卡控流程", ASSET_DIR / "03-wait-pilot-control.png",
               "根据 WatchDog、Pilot 所在站点和 Other Lot 当前站点判断是否增加或解除 WaitPilotChangeFOUP 卡控。")

    add_heading(doc, "3.2 LithoRule 新增卡控逻辑", 3)
    add_heading(doc, "3.2.1 R2RAutoPirunControl", 4)
    add_body(doc, "针对非 Specify Lot，如果同时满足以下条件，则增加 `Reason=R2RAutoPirunControl`；否则不增加该卡控：", keep=True)
    add_code_block(doc, """Pi_split_flag='Y'
AND (
    R2R CD Status='PIRUNON'
    OR R2R OVL Status='PIRUNON'
)
AND Pilot IS NOT NULL""")

    add_heading(doc, "3.2.2 Parent&ChildLotNeedRunSameTool", 4)
    add_heading(doc, "3.2.2.1 判断 Lot 是否存在 Pretool", 4)
    add_body(doc, "从 `r2r_litho_context_relation` 获取 `productid`、`curr_layer`、`pre_layer`，从 `r2r_litho_context_ovl` 获取 `productid`、`layerid`、`pretool`。按 `productid + pre_layer` 获取 `curr_layer`，再按 `productid + curr_layer` 获取 `pretool`。")
    add_body(doc, "只有 `pretool` 不为空时，才继续判断子母批同机台；`pretool` 为空时，不执行该卡控。")

    add_heading(doc, "3.2.2.2 获取子母批作业机台", 4)
    add_body(doc, "通过 `fabfutureaction` 获取与当前 Lot 存在 FutureMerge 关系的子批或母批 Lot。子母批关系范围以现有 `FutureMerge` 返回结果为准，不额外限定直接关系或完整关系链。")
    add_body(doc, "从 `r2r_lot_history` 按 `lotid + productid + layerid` 获取子批或母批在待判断 Lot 当前 Layer 的作业机台 `toolid`。最新一笔记录按照实际作业完成时间降序、记录 ID 降序获取。")

    add_heading(doc, "3.2.2.3 判断是否需要卡控", 4)
    add_body(doc, "通过 `r2r_litho_whitelist` 按 `productid`、`layerid`、`lotid` 判断是否为 Specify Lot。针对非 Specify Lot，如果待判断 Lot 的机台与子批或母批的作业机台 `toolid` 不一致，则增加 `Reason=Parent&ChildLotNeedRunSameTool`；否则按现有逻辑继续判断。")

    add_heading(doc, "4. LithoAssign 新增卡控逻辑", 2)
    add_body(doc, "LithoAssign 增加子母批使用同一机台的卡控，同时保留原有 `R2RAutoPirunControl` 逻辑。")
    add_heading(doc, "4.1 R2RAutoPirunControl", 3)
    add_body(doc, "针对非 Specify Lot，如果 `Pi_split_flag='Y'`，R2R CD 或 OVL Status 为 `PIRUNON`，且 Pilot 不为空，则增加 `Reason=R2RAutoPirunControl`；否则不增加该卡控。")
    add_heading(doc, "4.2 Parent&ChildLotNeedRunSameTool", 3)
    add_body(doc, "判断逻辑与 LithoRule 一致。`pretool` 及子批、母批作业机台信息改由 Central 获取，其他判断规则保持一致。")

    add_heading(doc, "二、AMA 新增逻辑", 2)
    add_body(doc, "AMA 根据 Report `Central_GetLithoR2RAutoPirunInfo` 执行 Pilot 设置和物理分批，并根据 Report `LithoPilotAutoDoAdhocSorter` 执行 Transfer FOUP。")

    add_heading(doc, "1. 设置 Pilot", 2)
    add_body(doc, "从 `Central_GetLithoR2RAutoPirunInfo` 获取 Lot、`toolid`、`productid`、`layerid`、`reticleid`、`prereticle`、`pretool`、`custom_context_value`、`pi_splitwafer`、`IsNeedSplit`、`isSTNSite`。")
    add_heading(doc, "1.1 整批设置为 Pilot", 3)
    add_body(doc, "当 `IsNeedSplit=F` 时，直接将整批 Lot 传给 R2R。")

    add_heading(doc, "1.2 物理分批 Pilot", 3)
    add_body(doc, "当 `IsNeedSplit=T` 时，执行以下处理。")
    add_heading(doc, "1.2.1 执行前复核", 4)
    add_numbers(doc, [
        "Lot 属于当前执行工厂，即 FAB6 或 FAB8；判断方式沿用现有程序逻辑。",
        "`fwlot.extrastatus='WaitForJobPrep'`。",
        "Lot 当前站点的 `runcardid` 为空。",
        "Lot 当前 Capability 为 `LithoCapability`、`L-BARCO-L` 或 `L-BARCO-S`。",
        "`CarrierKind='FOUP'`。",
        "Report 中选中的 Wafer 当前仍归属于该 Lot；仅检查归属关系，不额外检查 Wafer 状态、Slot 或 Chuck 信息。",
    ])
    add_body(doc, "任一条件不满足时，停止处理该 Lot，记录具体失败原因，等待下一轮重新计算，不回退为整批 Pilot。本需求不新增重复分批防护机制，由现有系统保证同一 Lot 不会被重复物理分批。")

    add_heading(doc, "1.2.2 预占空 FOUP", 4)
    add_body(doc, "执行前复核通过后，先获取并预占一个空 FOUP，再调用 MES 物理分批接口。如果未获取到可用空 FOUP，则不执行物理分批，将整批 Lot 传给 R2R。")

    add_heading(doc, "1.2.3 调用 MES 物理分批接口", 4)
    add_body(doc, "调用 MES 物理分批接口，将 `pi_splitwafer` 从原 Lot 中分出。调用成功时，将分出的子批设置为 Pilot，并将子批 Pilot 传给 R2R。")
    add_body(doc, "调用失败时，立即释放已预占的空 FOUP，并将整批 Lot 作为 Pilot 传给 R2R。物理分批失败并回退整批时，`IsNeedSplit`、`pi_splitwafer` 等字段的后续处理沿用现有 AMA 逻辑。向 R2R 设置 Pilot 与 Transfer FOUP 的先后顺序同样沿用现有 AMA 流程。")

    add_figure(doc, "图 4  AMA 物理分批与回退流程", ASSET_DIR / "04-ama-split.png",
               "AMA 根据 IsNeedSplit 判断整批或物理分批，物理分批前完成六项复核、预占 FOUP，并按接口结果设置子批或回退整批。",
               width_inches=4.55)

    add_heading(doc, "2. Transfer FOUP", 2)
    add_body(doc, "从 Report `LithoPilotAutoDoAdhocSorter` 获取待处理记录。AMA 按 Report 顺序调用 MES Transfer FOUP 接口，将 Pilot 导入已获取的空 FOUP。")
    add_body(doc, "如果无法获取可用空 FOUP、可用空 FOUP 数量为 0，或 MES Transfer FOUP 接口调用失败，则在 `AMALog` 中记录失败信息。Other Lot 紧急解除卡控后，Pilot 的 Transfer FOUP 任务由现有 Adhoc Sorter 流程处理。")

    add_heading(doc, "四、测试及验收说明", 1)
    add_body(doc, "本需求单不展开 Test Case 和验收场景。整批 Pilot、物理分批、跨厂、无空 FOUP、MES 接口失败、BARCO、Transfer FOUP、子母批同机台等验收场景，另行建立独立测试文档。")

    approval_heading = add_heading(doc, "五、审批意见", 1)
    approval_heading.paragraph_format.page_break_before = True
    add_approval_table(doc)

    # Core properties
    doc.core_properties.title = "LithoAutoPiRun 新增需求申请单（修订稿）"
    doc.core_properties.subject = "逻辑分批改造为物理分批及 Pilot 选择、Transfer FOUP、卡控逻辑优化"
    doc.core_properties.author = "制造部"
    doc.core_properties.keywords = "LithoAutoPiRun, RTD, AMA, Pilot, FOUP, 物理分批"

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
