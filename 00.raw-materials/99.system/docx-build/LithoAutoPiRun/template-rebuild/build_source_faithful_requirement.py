from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(r"D:\Obsidian\work\OBSidianCodex")
BUILD_DIR = ROOT / "00.raw-materials" / "99.system" / "docx-build" / "LithoAutoPiRun"
ASSET_DIR = BUILD_DIR / "assets"
OUT_DIR = ROOT / "00.raw-materials" / "90.processed" / "LithoAutoPiRun"
OUT_PATH = OUT_DIR / "LithoAutoPiRun_需求申请单_最终版.docx"
QA_DIR = BUILD_DIR / "source-faithful-qa"
QA_PATH = QA_DIR / "source-faithful-requirement.docx"

BLACK = "000000"


def set_run_font(run, size=10.5, bold=False, italic=False,
                 ascii_font="Times New Roman", east_asia="SimSun"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    for shading in list(rpr.findall(qn("w:shd"))):
        rpr.remove(shading)


def set_style_font(style, size=10.5, bold=False,
                   ascii_font="Times New Roman", east_asia="SimSun"):
    style.font.name = ascii_font
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.bold = bold


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, size=10):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), BLACK)


def set_table_geometry(table, total_width_twips=10380):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    column_count = len(table.columns)
    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    base_width = total_width_twips // column_count
    for _ in range(column_count):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(base_width))
        grid.append(col)
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(base_width))
            tc_w.set(qn("w:type"), "dxa")


def set_cant_split(row, enabled=True):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if enabled and existing is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    elif not enabled and existing is not None:
        tr_pr.remove(existing)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_runs(paragraph, parts, size=10.5):
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)


def add_cell_paragraph(cell, text="", *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_line=True, before=0, after=2, keep=False, size=10.5):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(cell, text, level=1):
    before = {1: 3, 2: 3, 3: 2, 4: 1}.get(level, 1)
    p = add_cell_paragraph(
        cell,
        text,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
        before=before,
        after=1,
        keep=True,
    )
    return p


def add_item(cell, text, indent_chars=2):
    p = add_cell_paragraph(
        cell,
        text,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line=False,
        after=1,
    )
    p.paragraph_format.left_indent = Pt(10.5 * indent_chars)
    p.paragraph_format.hanging_indent = Pt(0)
    return p


def add_flowchart(cell, caption, image_name, width_inches=5.75):
    figure_table = cell.add_table(rows=1, cols=1)
    figure_table.autofit = False
    figure_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = figure_table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tbl_pr.append(borders)
    set_cant_split(figure_table.rows[0])
    figure_cell = figure_table.cell(0, 0)
    set_cell_margins(figure_cell, top=30, start=0, bottom=30, end=0)
    p = clear_cell(figure_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption)
    set_run_font(run, size=10.5, bold=True)
    p2 = figure_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.0
    run = p2.add_run()
    run.add_picture(str(ASSET_DIR / image_name), width=Inches(width_inches))


def add_content(cell):
    clear_cell(cell)
    add_heading(cell, "需求内容（可添加附件）：", 1)
    add_heading(cell, "■　方案逻辑：", 1)

    add_heading(cell, "一、RTD新增逻辑", 1)
    add_cell_paragraph(
        cell,
        "1. 修改 Report：Central_GetLithoR2RAutoPirunInfo 中的 Pilot 选择逻辑；2. 增加 Report：LithoPiLotAutoDoAdhocSorter；3. Rule 和 Assign 中增加子母批作业同一机台逻辑。",
    )

    add_heading(cell, "1.Report：Central_GetLithoR2RAutoPirunInfo部分", 1)
    add_cell_paragraph(cell, "对选 Lot 和选片逻辑进行修改。")

    add_heading(cell, "1.1 Lot获取", 2)
    add_cell_paragraph(cell, "从 FAB6 和 FAB8 两厂获取待判断 Lot，并进行初步过滤。")

    add_heading(cell, "1.1.1 Lot基础信息获取", 3)
    add_cell_paragraph(cell, "通过表 fwlot 获取 appid、priority、processingstatus、componentqty 栏位信息；")
    add_cell_paragraph(cell, "通过表 fablotext 获取 requiredcapability、runcardid、reticleid 栏位信息；")
    add_cell_paragraph(cell, "通过表 fablotcarrierext 获取 carrierkind 栏位信息；")
    add_cell_paragraph(cell, "通过 fabinqtimeprocess 获取 RemainQ 栏位信息；")
    add_cell_paragraph(cell, "从 UI.RTDConfig-LITHOLotAssignment-LithoAssignCapability 中获取 LithoCapability。BARCO Capability 为 L-BARCO-L、L-BARCO-S。")

    add_heading(cell, "1.1.2 基础过滤条件", 3)
    add_cell_paragraph(cell, "对拿取的 Lot 执行以下过滤：")
    add_item(cell, "（1）基本条件过滤：筛选同时满足 processingstatus='Active' 或 'CrossFabTransferred'、carrierkind='FOUP'、runcardid 为空、requiredcapability 为 LithoCapability、L-BARCO-L 或 L-BARCO-S 的 Lot。")
    add_item(cell, "（2）跨厂信息过滤：筛选满足 IsTransferLot=True，或 IsTransferLot≠True 且 processingstatus≠'CrossFabTransferred' 的 Lot。IsTransferLot 指标沿用 istransferlot Macro 判断；本需求不新增 FAB6/FAB8 跨厂资料去重规则。")

    add_heading(cell, "1.2 获取Pirun站点并对Lot进一步过滤", 2)
    add_cell_paragraph(cell, "获取 Lot 后续站点信息并进一步判断是否符合 AutoPirun 条件。")

    add_heading(cell, "1.2.1 PirunLoop信息获取", 3)
    add_cell_paragraph(cell, "对通过 1.1 筛选的 Lot，向下 Fetch 20 个站点，获取每个站点的 productname、planname、stage、capability、stepseq、recipeid、STN 信息（Transfer Lot 沿用现有厂别修正规则）。将 Fetch 的站点截取到相同 Stage 最后一道 CD 站点，Lot 当前站点到最后一道 CD 站点之间的区段定义为 PirunLoop。若 Fetch 站点中无 CD 站点，则过滤该 Lot。")

    add_heading(cell, "1.2.2 其他过滤条件", 3)
    add_item(cell, "（1）Litho站点判断：PirunLoop 中必须存在 capability=LithoCapability 且含有 Reticle 信息的站点；若无则过滤 Lot。")
    add_item(cell, "（2）Specify判断：按 productid、layer、lotid 匹配 r2r_litho_whitelist；命中时沿用现有 Specify Lot 处理。")
    add_item(cell, "（3）同FOUP Pilot判断：从 rtd_r2r_litho_context_ovlcd 和 rtd_r2r_lot_history 获取 Litho Pilot，若同 FOUP 中已有 Litho Pilot，则过滤 Lot。")
    add_item(cell, "（4）FutureHold判断：从 fabfutureaction 获取 PirunLoop 各站点的 FutureAction；若存在 FutureHold，则过滤 Lot。")
    add_item(cell, "（5）RC判断：若 PirunLoop 中存在 RC 站点，则过滤 Lot。")

    add_heading(cell, "1.3 R2R条件判断", 2)
    add_cell_paragraph(cell, "获取 Lot 在可作业机台的 R2R 状态，并判断 Lot 是否存在多路径。")

    add_heading(cell, "1.3.1 可作业机台获取", 3)
    add_cell_paragraph(cell, "获取 1.2 中 Lot 的 Litho 站点。经现有 Transfer Macro 判断后，若 Lot 在 Litho 站点能够 Transfer，则同时获取对应工厂的 Litho 机台。")
    add_cell_paragraph(cell, "按机台厂别从 rtd_r2r_litho_add_setting 获取 Pi_split_flag 和 pi_splitcnt。若 Lot 不存在 Pi_split_flag='Y' 的可作业机台，则过滤 Lot。")
    add_cell_paragraph(cell, "剩余机台继续经过 EQPStatus、LCC、Capability、Recipe、PPID、Global Reason 判断；若存在卡控，则剔除对应机台。")

    add_heading(cell, "1.3.2 多路径判断", 3)
    add_cell_paragraph(cell, "按 Lot+STN+Reticle 维度从 rtd_r2r_litho_context_ovl 和 rtd_r2r_litho_context_cd 匹配 R2R 状态。保留 OVL_Status、CD_Status 属于 PIRUNON、ON、Fixed，且不存在 R2R Reason 的 Context。")
    add_cell_paragraph(cell, "按 Lot 统计有效 Context 数量。若 ContextCount>1，则判定 Lot 存在多路径并过滤该 Lot。")

    add_heading(cell, "1.4 选Lot规则", 2)
    add_cell_paragraph(cell, "按 Context 为 Lot 排序并挑选最优 Lot。")

    add_heading(cell, "1.4.1 AutoPirun Context筛选", 3)
    add_cell_paragraph(cell, "获取 1.3 过滤后的 Lot 及 Context 信息，栏位包括 Lot、STN、Reticle、Prod、Layer、Recipe、Pretool、Prereticle、Custom_Context_Value、CD_Status、OVL_Status、Pilot_CD、Pilot_OVL、Pi_split_flag、pi_splitcnt。")
    add_cell_paragraph(cell, "筛选 Pi_split_flag='Y'，且 Pilot_CD 为空或 Pilot_OVL 为空的 Context，作为需要自动 Pirun 的 Context。")

    add_heading(cell, "1.4.2 Context内Lot排序", 3)
    add_cell_paragraph(cell, "获取 Lot 的排序指标：")
    add_item(cell, "（1）GapToLitho：Lot 当前站点距 Litho 站点的剩余 Step 数量。")
    add_item(cell, "（2）SplitCntMatched：componentqty≥pi_splitcnt 时为 1，否则为 0。pi_splitcnt 为空、等于 0、小于 0 或大于 25 时使用默认值 4。")
    add_item(cell, "（3）RequiredChuckCount：Lot 存在 prelayer 时，从 r2r_litho_waferhistory 获取 Wafer 的 Chuck 信息；否则从 fsmaterialassociation 获取 Wafer 的 Slot 信息。C1、C2 各至少 2 片，或奇数、偶数 Slot 各至少 2 片时为 1，否则为 0。")
    add_item(cell, "（4）BulletLot：沿用现有逻辑判断 Lot 是否为空扣或空 LP Lot；是为 1，否则为 0。")
    add_item(cell, "（5）RemainQ：RemainQ 有值且大于 0 时取 Lot 实际剩余 Queue Time；空值按 9999 h。")
    add_item(cell, "（6）KeyLot：quota_applyinfo 中 KeyLot=1 且 Status='CONFIRM' 时为 1，否则为 0。")
    add_cell_paragraph(cell, "按以下优先级对 Lot 排序，记为 RTDRank：")
    for text in (
        "① GapToLitho ASC",
        "② SplitCntMatched DESC",
        "③ RequiredChuckCount DESC",
        "④ BulletLot DESC",
        "⑤ RemainQ ASC",
        "⑥ KeyLot ASC",
        "⑦ lotid ASC（最终 Tie-breaker）",
    ):
        add_item(cell, text, indent_chars=3)

    add_heading(cell, "1.4.3 Context之间排序", 3)
    add_cell_paragraph(cell, "Lot 存在多个 Pirun Context 时，为尽量将使用同一 Reticle 的 Lot 分配到同一机台，并使不同 Reticle 的 Lot 均衡分配，需要对 Context 进行排序。")
    add_item(cell, "（1）ReticleSTNRank：按 Reticle+STN 对 Context 分组并获取 ReticleOnSTN。第一轮优先 STN=ReticleOnSTN 的 Context；后续轮次优先与上一轮已选 Context 属于同一 Reticle+STN 分组的 Context。")
    add_item(cell, "（2）ContextCandidateCount：统计 Context 内当前可选 Lot 数量。")
    add_item(cell, "（3）ActualSTNPilotCount：按 STN 统计已经选择 Pilot 的 Context 数量。")
    add_cell_paragraph(cell, "最终按 ReticleSTNRank DESC、ContextCandidateCount ASC、ActualSTNPilotCount ASC、RTDRank ASC 对 Lot+Context 排序。")

    add_heading(cell, "1.4.4 Context循环挑选Pilot", 3)
    add_cell_paragraph(cell, "将排序第一的 Lot+Context 固定为已选 Pilot Context，并剔除与已选 Context 相同的其他 Lot+Context；更新 ReticleSTNRank、ContextCandidateCount、ActualSTNPilotCount 后进入下一轮，直至无可用 Context 或无可用 Lot。每个 Context 最多选择一个 Pilot。")
    add_flowchart(cell, "图1　RTD候选筛选与Pilot选择流程", "01-rtd-selection.png")

    add_heading(cell, "1.5 Pilot选片逻辑", 2)
    add_cell_paragraph(cell, "判断 Lot 是否需要整批设为 Pilot；不满足整批条件时，按规则选片进行物理分批。")

    add_heading(cell, "1.5.1 整批设为Pilot场景", 3)
    add_cell_paragraph(cell, "以下五项为“或”关系，任一项成立时设置 IsNeedSplit=F，并将整批 Lot 设为 Pilot：")
    for text in (
        "（1）BulletLot=1 或 KeyLot=1；",
        "（2）CurCapability=LithoCapability 且 FuLL(RemainQ)；",
        "（3）RequiredChuckCount=0；",
        "（4）SplitCntMatched=0；",
        "（5）componentqty≤6。",
    ):
        add_item(cell, text)
    add_cell_paragraph(cell, "componentqty>0 由上游系统保证，本需求不增加零片判断。")

    add_heading(cell, "1.5.2 物理分批选片逻辑", 3)
    add_cell_paragraph(cell, "不满足 1.5.1 任一条件时，设置 IsNeedSplit=T，并执行物理分批选片。")

    add_heading(cell, "1.5.2.1 Wafer分组", 4)
    add_item(cell, "（1）Group层：Wafer ID 1～10 划入 Group1，Wafer ID 11～25 划入 Group2。")
    add_item(cell, "（2）SubGroup层：有 Chuck 信息时，C1 划入 SubGroup1，C2 划入 SubGroup2；无 Chuck 信息时，奇数 Slot 划入 SubGroup1，偶数 Slot 划入 SubGroup2。")
    add_item(cell, "（3）GroupRank：Group1+SubGroup1=1；Group1+SubGroup2=2；Group2+SubGroup1=3；Group2+SubGroup2=4。")

    add_heading(cell, "1.5.2.2 Wafer排序与选择", 4)
    add_cell_paragraph(cell, "在 Group+SubGroup 内按 Wafer ID 升序编号为 WaferRank，最终按 WaferRank ASC、GroupRank ASC 轮流选片。")
    add_cell_paragraph(cell, "pi_splitcnt 为空、等于 0、小于 0 或大于 25 时，使用默认值 4；有效范围为 1～25。若选片数大于 Lot 当前可用 Wafer 数，则不执行物理分批，改为整批 Pilot。选中的 Wafer 记为 pi_splitwafer，物理分批后的子批设置为 Pilot。")

    add_heading(cell, "1.5.3 Merge站点设定", 3)
    add_cell_paragraph(cell, "若 Lot 存在不包含 SRC 的 ADI 站点，则将第一道符合条件的 ADI 设置为 Merge 站点；否则将最后一道 CD 设置为 Merge 站点。有效 Merge 站点由现有流程保证。")
    add_flowchart(cell, "图2　Pilot整批与物理分批判断流程", "02-pilot-split.png")

    add_heading(cell, "1.6 输出Report", 2)
    add_cell_paragraph(cell, "将计算结果存入 Report：Central_GetLithoR2RAutoPirunInfo，供 AMA 读取执行。")
    add_cell_paragraph(cell, "Report 栏位包括：Lot、toolid、productid、layerid、reticleid、prereticle、pretool、custom_context_value、pi_splitwafer、IsNeedSplit、isSTNSite。")
    add_cell_paragraph(cell, "保持现有 Report 输出范围和记录方式，不新增 request_id、generated_time、execution_status、error_code。")

    add_heading(cell, "2.新增Report：LithoPiLotAutoDoAdhocSorter", 1)
    add_cell_paragraph(cell, "获取需要 TransferFoup 的 Pilot。")

    add_heading(cell, "2.1 Litho Pilot拿取", 2)
    add_cell_paragraph(cell, "获取需要 TransferFoup 的 Litho Pilot，判断逻辑如下：")
    add_heading(cell, "2.1.1 数据获取", 3)
    add_cell_paragraph(cell, "从 r2r_litho_context_ovl 获取 ovl_status、pilot；从 r2r_litho_context_cd 获取 cd_status、pilot；从 fwlot 获取 appid、lottype、priority；从 fabcategorymap 获取 lottype、category。")
    add_heading(cell, "2.1.2 筛选判断", 3)
    add_cell_paragraph(cell, "筛选 priority<5、category='Production'，且 ovl_status='PIRUNON' 或 cd_status='PIRUNON' 的 Pilot，作为需要 TransferFoup 的 Litho Pilot。")

    add_heading(cell, "2.2 Transfer FOUP判断", 2)
    add_cell_paragraph(cell, "获取 2.1 中 Litho Pilot 的 Carrier 信息，按 Carrier 从 fwlot 获取所有 Lot，并保留 extrastatus='WaitForJobPrep' 的记录。若 Litho Pilot 同 Carrier 中存在其他 Lot，则该 Pilot 需要 Change Foup。")

    add_heading(cell, "2.3 Carrier排序规则", 2)
    add_cell_paragraph(cell, "对需要 TransferFOUP 的 Litho Pilot，按 RemainQ ASC、Priority ASC、componentqty DESC、lotid ASC 排序，并根据排序建立 AdhocSorterJob。")

    add_heading(cell, "2.4 输出Report", 2)
    add_cell_paragraph(cell, "从 AMA.TriggerConfig-WatchDog_LithoPiLotAutoDoAdhocSorter 获取 Switch、Trigger Time Slot、TriggerCount/Time。当 Switch='Y' 且当前时间位于现有 Trigger Time Slot 范围内时，按 TriggerCount/Time 选择 Carrier，并将结果存入 Report：LithoPiLotAutoDoAdhocSorter。栏位包括 Carrier、Pilot、extrastatus、Status、RemainQ、Pieces、Prod、Priority。")

    add_heading(cell, "3.Rule中新增卡控逻辑", 1)
    add_cell_paragraph(cell, "在 Global Macro 中新增需要 Transfer Foup 的 Litho Pilot 卡控；在 LithoRule 中增加子母批 Run 相同机台的卡控。")

    add_heading(cell, "3.1 Global Macro新增卡控逻辑", 2)
    add_cell_paragraph(cell, "拿取 2.1 中的 Litho Pilot。当 WatchDog_LithoPiLotAutoDoAdhocSorter 的 Switch='Y' 且当前时间位于 Trigger Time Slot 范围内时，按以下场景处理：")
    add_item(cell, "（1）若 Litho Pilot 位于 AdhocSorter 站点（adhocplanname 包含 UnscheduledSorter），则同 FOUP 中不在 AdhocSorter 站点的 Other Lot 增加 Reason=WaitPilotChangeFOUP。Other Lot 位于 Litho 站点时不解除；位于 BARCO 站点时，RemainQ<4 h 或触发 Qu_0 即解除；位于其他站点时，仅触发 Qu_0 时解除。")
    add_item(cell, "（2）若 Litho Pilot 不在 AdhocSorter 站点，则同 FOUP 中不在 AdhocSorter 站点的 Other Lot 增加 Reason=WaitPilotChangeFOUP。Other Lot 的解除规则同（1）；当 Litho Pilot 的 RemainQ<4 h 或触发 Qu_0 时，解除相关卡控。")
    add_flowchart(cell, "图3　WaitPilotChangeFOUP卡控流程", "03-wait-pilot-control.png")

    add_heading(cell, "3.2 Litho Rule新增卡控逻辑", 2)
    add_heading(cell, "3.2.1 R2RAutoPirunControl卡控", 3)
    add_cell_paragraph(cell, "针对非 Specify Lot，若 Pi_split_flag='Y'，R2R CD 或 OVL Status='PIRUNON'，且 Pilot 不为空，则增加 Reason=R2RAutoPirunControl；否则不增加卡控。")

    add_heading(cell, "3.2.2 Parent&ChildLotNeedRunSameTool卡控", 3)
    add_heading(cell, "3.2.2.1 判断Lot是否有Pretool", 4)
    add_cell_paragraph(cell, "从 r2r_litho_context_relation 获取 productid、curr_layer、pre_layer；从 r2r_litho_context_ovl 获取 productid、layerid、pretool。按 productid+pre_layer 获取 curr_layer，再按 productid+curr_layer 获取 pretool；仅当 pretool 不为空时继续判断。")
    add_heading(cell, "3.2.2.2 获取子母批作业机台", 4)
    add_cell_paragraph(cell, "通过 fabfutureaction 获取与当前 Lot 存在 FutureMerge 关系的子批或母批 Lot；从 r2r_lot_history 按 lotid+productid+layerid 获取其在待判断 Lot 当前 Layer 的作业机台 toolid，并按实际作业完成时间、记录 ID 降序取最新记录。")
    add_heading(cell, "3.2.2.3 判断是否需卡控", 4)
    add_cell_paragraph(cell, "通过 r2r_litho_whitelist 按 productid、layerid、lotid 判断 Specify Lot。针对非 Specify Lot，若待判断 Lot 的机台与子批或母批作业机台 toolid 不一致，则增加 Reason=Parent&ChildLotNeedRunSameTool；否则按原逻辑判断。")

    add_heading(cell, "4.LithoAssign新增卡控逻辑", 1)
    add_cell_paragraph(cell, "LithoAssign 增加子母批 Run 相同机台的卡控，并保留原 R2RAutoPirunControl 逻辑。")
    add_heading(cell, "4.1 R2RAutoPirunControl卡控", 2)
    add_cell_paragraph(cell, "针对非 Specify Lot，若 Pi_split_flag='Y'，R2R CD 或 OVL Status='PIRUNON'，且 Pilot 不为空，则增加 Reason=R2RAutoPirunControl；否则不增加卡控。")
    add_heading(cell, "4.2 Parent&ChildLotNeedRunSameTool卡控", 2)
    add_cell_paragraph(cell, "逻辑与 LithoRule 一致，pretool 及子批、母批作业机台信息改由 Central 获取。")

    add_heading(cell, "二、AMA新增逻辑", 1)
    add_cell_paragraph(cell, "根据 Report：Central_GetLithoR2RAutoPirunInfo 执行 Pilot 设置和物理分批，并将 Pilot 传给 R2R；根据 Report：LithoPiLotAutoDoAdhocSorter 执行 TransferFoup。")

    add_heading(cell, "1. 设置Pilot", 2)
    add_cell_paragraph(cell, "获取 Report：Central_GetLithoR2RAutoPirunInfo 栏位信息：Lot、toolid、productid、layerid、reticleid、prereticle、pretool、custom_context_value、pi_splitwafer、IsNeedSplit、isSTNSite。")
    add_heading(cell, "1.1 整批设为Pilot", 3)
    add_cell_paragraph(cell, "当 IsNeedSplit=F 时，直接将整批 Lot 传给 R2R。")

    add_heading(cell, "1.2 物理分批Pilot", 3)
    add_cell_paragraph(cell, "当 IsNeedSplit=T 时，按以下顺序处理。")
    add_heading(cell, "1.2.1 执行前复核", 4)
    for text in (
        "（1）分批 Lot 属于当前执行工厂，即 FAB6 或 FAB8；",
        "（2）fwlot.extrastatus='WaitForJobPrep'；",
        "（3）Lot 当前站点的 runcardid 为空；",
        "（4）Lot 当前 Capability 为 LithoCapability、L-BARCO-L 或 L-BARCO-S；",
        "（5）CarrierKind='FOUP'；",
        "（6）Report 中选中的 Wafer 当前仍归属于该 Lot。",
    ):
        add_item(cell, text)
    add_cell_paragraph(cell, "仅检查 Wafer 归属关系，不额外检查 Wafer 状态、Slot 或 Chuck。任一复核项不满足时，停止处理该 Lot，记录失败原因并等待下一轮重新计算，不回退为整批 Pilot。重复物理分批防护沿用现有系统逻辑。")

    add_heading(cell, "1.2.2 空FOUP与MES物理分批", 4)
    add_cell_paragraph(cell, "六项复核通过后，先获取并预占空 FOUP，再调用 MES 物理分批接口。未取得空 FOUP 时，不执行物理分批，将整批 Lot 传给 R2R。")
    add_cell_paragraph(cell, "接口成功时，将 pi_splitwafer 从原 Lot 中分出，生成子批并设置为 Pilot，再按现有顺序传给 R2R、执行 Transfer FOUP。接口失败时，立即释放预占 FOUP，将整批 Lot 作为 Pilot 传给 R2R；IsNeedSplit、pi_splitwafer 的后续处理沿用现有 AMA 逻辑。")
    add_flowchart(cell, "图4　AMA物理分批与回退流程", "04-ama-split.png")

    add_heading(cell, "2. TransferFoup", 2)
    add_cell_paragraph(cell, "获取 Report：LithoPiLotAutoDoAdhocSorter 栏位信息，并按顺序调用 MES Transfer FOUP 接口，将 Pilot 导入空 FOUP。若无法取得可用空 FOUP、空 FOUP 数量为 0，或 MES Transfer FOUP 接口失败，则在 AMALog 中记录失败信息；后续继续沿用现有 Adhoc Sorter 流程。")


def add_content_exact_replica(cell):
    """Reproduce the source-image wording, with only confirmed Q&A changes."""
    clear_cell(cell)
    add_heading(cell, "需求内容（可添加附件）：", 1)
    add_heading(cell, "■  方案逻辑：", 1)

    add_heading(cell, "一、 RTD新增逻辑", 1)
    add_cell_paragraph(cell, "1. 修改 Report： Central_GetLithoR2RAutoPirunInfo 中的 Pilot 选择逻辑；2.增加 Report：LithoPiLotAutoDoAdhocSorter；3.Rule和 Assign 中增加子母批作业同一机台逻辑")

    add_heading(cell, "1.Report：Central_GetLithoR2RAutoPirunInfo部分", 1)
    add_cell_paragraph(cell, "对选 lot 和选片逻辑进行修改。")

    add_heading(cell, "1.1 Lot获取", 2)
    add_cell_paragraph(cell, "从 FAB6 和 FAB8 两厂获取待判断 Lot，并进行初步过滤。")

    add_heading(cell, "1.1.1 Lot基础信息获取", 3)
    add_cell_paragraph(cell, "通过表fwlot获取appid、priority、processingstatus、componentqty栏位信息；")
    add_cell_paragraph(cell, "通过表fablotext获取requiredcapability、runcardid、reticleid栏位信息；")
    add_cell_paragraph(cell, "通过表fablotcarrierext获取carrierkind栏位信息；")
    add_cell_paragraph(cell, "通过fabinqtimeprocess获取RemainQ栏位信息；")
    add_cell_paragraph(cell, "从 UI.RTDConfig-LITHOLotAssignment-LithoAssignCapability 中获取 LithoCapability；BARCO Capability 固定为 L-BARCO-L、L-BARCO-S。")

    add_heading(cell, "1.1.2 基础过滤条件", 3)
    add_cell_paragraph(cell, "对拿取的 Lot 执行以下过滤：")
    add_item(cell, "（1）基本条件过滤：筛选出满足 processingstatus = 'Active' or 'CrossFabTransferred'，carrierkind = 'FOUP'、runcardid 为空、requiredcapability In（LithoCapability，L-BARCO-L，L-BARCO-S）的 Lot。")
    add_item(cell, "（2）跨厂信息去重：筛选出满足 IsTransferLot =True 或（IsTransferLot ≠ True 且 processingStatus ≠ 'CrossFabTransferred'）的 Lot。（IsTransferLot 指标经 istransferlot marco 判断得到）")

    add_heading(cell, "1.2 获取Pirun站点并对lot进一步过滤", 2)
    add_cell_paragraph(cell, "获取 Lot 后续站点信息并进一步判断是否符合AutoPirun条件。")

    add_heading(cell, "1.2.1 PirunLoop信息获取", 3)
    add_cell_paragraph(cell, "对通过 1.1 筛选的 Lot，向下 Fetch 20 站，获取每个站点的 productname、planname、stage、capability、stepseq、recipeid、STN 信息(Transferlot 需修正厂别)。将 Fetch 的站点截取到同 Stage 最后一道 CD 站点，Lot 当前站点到最后一道 CD 站点即为一段 PirunLoop。若 lot Fetch 站点中无 CD 站点，则过滤该 Lot。")

    add_heading(cell, "1.2.2 其他过滤条件", 3)
    add_item(cell, "（1）Litho 站点判断：判断 PirunLoop 中是否存在 capability =LithoCapability 且含有 Reticle 信息的站点。若无则过滤 Lot。")
    add_item(cell, "（2）Specify 判断：判断 lot 是否在 r2r_litho_whitelist（匹配 productid，layer，lotid）中，若在则沿用现有 Specify Lot 处理。")
    add_item(cell, "（3）同 Foup Pilot 判断：从表 rtd_r2r_litho_context_ovlcd 和 rtd_r2r_lot_history 获取 Litho Pilot，判断同 Foup 中是否有 Litho Pilot，若有则过滤 Lot。")
    add_item(cell, "（4）FutureHold 判断：从表 fabfutureaction 获取 PirunLoop 中每个站点的 FutureAction 信息，判断 Loop 中是否存在 FutureHold，若有则过滤 Lot。")
    add_item(cell, "（5）RC 判断：判断 PirunLoop 中是否存在 RC站点，若有则过滤 Lot。")

    add_heading(cell, "1.3 R2R条件判断", 2)
    add_cell_paragraph(cell, "获取lot在可作业机台的R2R 状态，并判断lot是否存在多路径。")

    add_heading(cell, "1.3.1 可作业机台获取", 3)
    add_cell_paragraph(cell, "拿取 1.2 中 Lot 的 Litho 站点，经 TransferMarco 判断后，若 Lot 在 Litho 站点能 Transfer 则同时拿取对厂 Litho 机台。")
    add_cell_paragraph(cell, "By 机台厂别从表 rtd_r2r_litho_add_setting 中获取 Lot 在机台的 Pi_split flag 和 pi_splitcnt。若 Lot 不存在 Pi_split flag = 'Y' 的机台，则过滤 Lot。")
    add_cell_paragraph(cell, "将剩余 lot 在 Litho 站点的机台经过 EQPStatus、LCC、Capability、Recipe、PPID、Global Reason 判断，若存在卡控则筛除对应机台。")

    add_heading(cell, "1.3.2 多路径判断", 3)
    add_cell_paragraph(cell, "By Lot+STN+Reticle 维度从表 rtd_r2r_litho_context_ovl 和 rtd_r2r_litho_context_cd 匹配 R2R 状态，并判断是否存在 R2R Reason。筛选出满足：OVL_Status In(PIRUNON，ON，Fixed)且 CD_Status In(PIRUNON，ON，Fixed)且无 R2R Reason 的 Context (By Lot+STN+Reticle)。")
    add_cell_paragraph(cell, "By Lot 统计符合上述条件的 Context 数量，若 ContextCount>1，则认为 Lot 存在多路径，过滤该 Lot。")

    add_heading(cell, "1.4 选Lot规则", 2)
    add_cell_paragraph(cell, "By Context 为 Lot 排序并挑选最优 Lot。")

    add_heading(cell, "1.4.1 AutoPirun Context筛选", 3)
    add_cell_paragraph(cell, "获取 1.3 过滤后的 Lot 及 Context 信息，栏位包括 Lot、STN、Reticle、Prod、Layer、Recipe、Pretool、Prereticle、Custom_Context_Value、CD_Status、OVL_Status、Pilot_CD、Pilot_OVL、Pi_split_flag、pi_splitcnt。")
    add_cell_paragraph(cell, "筛选出满足 Pi_split_flag='Y'且(Pilot_CD 为 Null 或 Pilot_OVL 为 Null)的 Context，即为需要自动 Pirun 的 Context。")

    add_heading(cell, "1.4.2 Context内Lot排序", 3)
    add_cell_paragraph(cell, "获取 lot 的排序指标")
    add_item(cell, "（1）计算 Lot 当前站点距 Litho 站点的剩余 Step 数量，记为 GapToLitho；")
    add_item(cell, "（2）判断 Lot 的片数是否大于等于 Pi_splitcnt（默认为4），若是则 SplitCntMatched=1，否则为0；Pi_splitcnt 为空、为0、为负数或大于25时使用默认值4；")
    add_item(cell, "（3）若 Lot 存在 prelayer，则从表 r2r_litho_waferhistory 中获取每片 Wafer 的 Chuck 信息；否则从表 fsmaterialassociation 中获取 Wafer 的 Slot 信息。判断 Lot 是否满足包含 C1/C2（或 Slot 奇/偶）各大于等于两片，若是则 RequiredChuckCount=1，否则为0；")
    add_item(cell, "（4）判断 Lot 是否为空扣/空 Lp Lot，若是则 BulletLot=1，否则为0；")
    add_item(cell, "（5）若 Lot RemainQ 有值且大于0，则指标 RemainQ 为 lot 剩余 Qtime，否则为9999；")
    add_item(cell, "（6）若 Lot 在表 quota_applyinfo 中且 KeyLot=1 且 Status=CONFIRM，则指标 KeyLot=1，否则为0。")
    add_cell_paragraph(cell, "按以下优先级对 Lot 排序，并记为 RTDRank：")
    for text in ("①Min(GapToLitho)", "②Max(SplitCntMatched)", "③Max(RequiredChuckCount)", "④Max(BulletLot)", "⑤Min(RemainQ)", "⑥Min(KeyLot)", "⑦Min(lotid)"):
        add_item(cell, text, indent_chars=3)

    add_heading(cell, "1.4.2 Context 之间排序", 3)
    add_cell_paragraph(cell, "Lot 存在多个 Pirun Context时，为实现使用同一 Reticle 的Lot尽量分配到同一机台，不同 Reticle 的 lot 均衡分配，需要对 Context 进行排序。")
    add_cell_paragraph(cell, "Context 排序指标：")
    add_item(cell, "（1）By Reticle+STN 给 Context 分组，获取 Reticle 当前所在机台信息 ReticleOnSTN，循环前，若 Context 的 STN=ReticleOnSTN，则 ReticleSTNRank=1；循环中，若 Context 与上轮排序第一的 Context 属于同一组，则 ReticleSTNRank=1，否则为0。")
    add_item(cell, "（2）统计 Context 内当前可选 Lot 的数量，记作 ContextCandidateCount。")
    add_item(cell, "（3）By STN 统计已选择 Pilot 的 Context 数量，记作 ActualSTNPilotCount。")
    add_cell_paragraph(cell, "最终按以下优先级对 Lot +Context 排序：")
    for text in ("① Max(ReticleSTNRank)", "② Min(ContextCandidateCount)", "③ Min(ActualSTNPilotCount)", "④ Min(RTDRank)"):
        add_item(cell, text, indent_chars=3)

    add_heading(cell, "1.4.3 Context 循环挑选 Pilot", 3)
    add_cell_paragraph(cell, "将排序第一的 Lot+Context 固定，作为已选 Pilot 的 Context，并去除 lot/Context 与已选 Context 相同的其他 Lot+Context，在更新 ReticleSTNRank、ContextCandidateCount、ActualSTNPilotCount 指标后，进入下一轮循环，直至无可用 Context 或无可用 Lot 后，结束循环。每个 Context 最多选择一个 Pilot。")
    add_flowchart(cell, "图1  RTD候选筛选与Pilot选择流程", "01-rtd-selection.png")

    add_heading(cell, "1.5 Pilot选片逻辑", 2)
    add_cell_paragraph(cell, "判断 Lot 是否要整批设为 Pilot，不满足的需要选片进行物理分批。")

    add_heading(cell, "1.5.1 整批设为 Pilot 场景", 3)
    for text in (
        "（1）Lot 生效 BulletLot=1 或 KeyLot=1；",
        "（2）Lot 的 CurCapability=LithoCapability 且 FuLL(RemainQ)；",
        "（3）Lot 的指标 RequiredChuckCount=0；",
        "（4）Lot 的指标 SplitCntMatched=0；",
        "（5）Lot 的片数 componentqty<=6",
    ):
        add_item(cell, text)
    add_cell_paragraph(cell, "以上五项为“或”关系，任一项成立时将整批设为 Pilot。")

    add_heading(cell, "1.5.2 物理分批选片逻辑", 3)
    add_cell_paragraph(cell, "不满足 1.5.1 场景时，Pilot 需要物理分批，Flag:IsNeedSplit 生效 T，并进行选片，规则如下：")

    add_heading(cell, "1.5.2.1 Wafer 分组", 4)
    add_cell_paragraph(cell, "将 Lot 的 Wafer 分为以下层级：")
    add_item(cell, "（1）Group 层：按 Waferid 编号，Waferid #1-#10 的 Wafer 划入 Group1；Waferid #11-#25 的 Wafer 划入 Group2。")
    add_item(cell, "（2）SubGroup 层：若 Lot 有 Chuck 信息，则 Chuck1（C1）的 Wafer → SubGroup1，Chuck2（C2）的 Wafer → SubGroup2；否则按 SlotMap，奇数 Slot 的 Wafer → SubGroup1，偶数 Slot 的 Wafer → SubGroup2。")
    add_item(cell, "（3）GroupRank 赋值：Group1-SubGroup1 =1；Group1-SubGroup2 =2；Group2-SubGroup1 =3；Group2-SubGroup2 =4。")

    add_heading(cell, "1.5.2.2 Wafer 排序与选择", 4)
    add_cell_paragraph(cell, "在 Group + SubGroup 内按 waferid 排序并编号，记为 WaferRank。按以下优先级排序：① MIN(WaferRank)；② MIN(GroupRank)。")
    add_cell_paragraph(cell, "选择规则：")
    add_cell_paragraph(cell, "若 Context 对应的 pi_splitcnt 有值，则按排序顺序挑选 pi_splitcnt 片 Wafer 作为 pi_splitwafer；pi_splitcnt 为空、为0、为负数或大于25时，使用默认值4。若选片数大于 Lot 当前可用 Wafer 数，则不执行物理分批，改为整批 Pilot。")
    add_cell_paragraph(cell, "将 pi_splitwafer 分批后设为 Pilot。")

    add_heading(cell, "1.5.3 Merge 站点设定", 3)
    add_cell_paragraph(cell, "若 Lot 有 ADI 站点（不包含 SRC），则 lot 需物理分批时，将第一道 ADI 设为 Merge 站点；否则将最后一道 CD 设为 Merge 站点。")
    add_flowchart(cell, "图2  Pilot整批与物理分批判断流程", "02-pilot-split.png")

    add_heading(cell, "1.6 输出 Report", 2)
    add_cell_paragraph(cell, "以上全部计算结果存入 Report：Central_GetLithoR2RAutoPirunInfo，供 AMA 读取执行。")
    add_cell_paragraph(cell, "Report 栏位包括：Lot、toolid、productid、layerid、reticleid、prereticle、pretool、custom_context_value、pi_splitwafer、IsNeedSplit、isSTNSite。")

    add_heading(cell, "2.新增Report：LithoPiLotAutoDoAdhocSorter", 1)
    add_cell_paragraph(cell, "获取需要 TransferFoup 的 Pilot。")

    add_heading(cell, "2.1  Litho Pilot拿取", 2)
    add_cell_paragraph(cell, "获取需要 TransferFoup 的 Litho Pilot，判断逻辑如下：")
    add_heading(cell, "2.1.1 数据获取：", 3)
    add_cell_paragraph(cell, "从 r2r_litho_context_ovl 获取 ovl_status、pilot 栏位信息；从 r2r_litho_context_cd 获取 cd_status、pilot 栏位信息；从 fwlot 中获取 appid、lottype、priority 栏位信息；从 fabcategorymap 中获取 lottype、category 栏位信息。")
    add_heading(cell, "2.1.2 筛选判断：", 3)
    add_cell_paragraph(cell, "筛选出满足 priority<5 且 category='Production' 且（ovl_status='PIRUNON' 或 cd_status='PIRUNON'）的 Pilot，即为需要 TransferFoup 的 Litho Pilot。")

    add_heading(cell, "2.2 Transfer FOUP判断", 2)
    add_cell_paragraph(cell, "拿取 2.1 中 LithoPilot 的 Carrier 信息，By Carrier 从 fwlot 中获取所有 Lot，保留 extrastatus='WaitForJobPrep' 的 Lot。")
    add_cell_paragraph(cell, "若 LithoPilot 同 Carrier 中存在其他 Lot 时，则该 Pilot 需要 Change Foup。")

    add_heading(cell, "2.3 Carrier排序规则", 2)
    add_cell_paragraph(cell, "对于需 TransferFOUP 的 LithoPilot 进行排序，拿取 LithoPilot 的 RemainQ、Priority、componentqty 指标，并按照 Min(RemainQ)、Min(Priority)、Max(componentqty)、Min(lotid)排序，根据排序建立 AdhocSorterJob。")

    add_heading(cell, "2.4输出Report", 2)
    add_cell_paragraph(cell, "从 AMA.TriggerConfig-WatchDog_LithoPiLotAutoDoAdhocSorter 拿取 Switch、Trigger Time Slot、TriggerCount/Time 栏位信息，当 Switch='Y' 且当前时间在 Trigger Time Slot 范围内时，将前 TriggerCount/Time 个需物理分批的 Carrier 结果存入 Report:LithoPiLotAutoDoAdhocSorter 中，栏位包括：Carrier、Pilot、extrastatus、Status、RemainQ、Pieces、Prod、Priority。")

    add_heading(cell, "3.Rule中新增卡控逻辑", 1)
    add_cell_paragraph(cell, "在 Global Macro 中新增对需导 Foup LithoPilot 的卡控；在 LithoRule 中增加子母批 Run 相同机台的卡控。")

    add_heading(cell, "3.1 Global Macro新增卡控逻辑", 2)
    add_cell_paragraph(cell, "拿取 2.1 中的 LithoPilot，当 WatchDog_LithoPiLotAutoDoAdhocSorter 中的 Switch='Y' 且当前时间在 Trigger Time Slot 范围内时，满足以下两种场景的 lot 需卡控 Reason：")
    add_item(cell, "（1）若 LithoPilot 在 AdhocSoter 站点（adhocplanname 包含 'UnScheduleSorter'），则 LithoPilot 同 Foup 不在 AdhocSoter 站点的 Other Lot 需要卡控 Reason: WaitPilotChangeFOUP。Remove 规则：当 Other Lot 在 Litho 站点时，不能 Remove；当 Other Lot 在 Barco 站点时，RemainQ<4H 或触发 Qu_0 时，Remove 卡控；当 Other Lot 在非 Litho/Barco 站点时，仅触发 Qu_0 时，Remove 卡控。")
    add_item(cell, "（2）若 LithoPilot 不在 AdhocSoter 站点，则 LithoPilot 和同 Foup 不在 AdhocSoter 站点的 Other Lot 都需要卡控 Reason: WaitPilotChangeFOUP。Remove 规则：Other Lot Follow（1），LithoPilot 的 RemainQ<4H 或触发 Qu_0 时，Remove 卡控。")
    add_flowchart(cell, "图3  WaitPilotChangeFOUP卡控流程", "03-wait-pilot-control.png")

    add_heading(cell, "3.2 Litho Rule 新增卡控逻辑", 2)
    add_heading(cell, "3.2.1 R2RAutoPirunControl卡控", 3)
    add_cell_paragraph(cell, "针对非Specify Lot，若Pi_SplitFlag='Y'且R2R CD/OVL Status=PIRUNON且Pilot不为Null，则卡控 Reason:R2RAutoPirunControl，否则不卡控。")

    add_heading(cell, "3.2.2 Parent&ChildLotNeedRunSameTool卡控", 3)
    add_heading(cell, "3.2.2.1 判断lot是否有Pretool", 4)
    add_cell_paragraph(cell, "从表r2r_litho_context_relation中获取productid、curr_layer、pre_layer栏位信息；从r2r_litho_context_ovl中获取productid、layerid、pretool栏位信息；")
    add_cell_paragraph(cell, "By Prod、layer从r2r_litho_context_relation（匹配productid、pre_layer）获取curr_layer，再通过Prod、curr_layer从r2r_litho_context_ovl获取pretool，当lot存在pretool不为空时，则需要后续判断。")
    add_heading(cell, "3.2.2.2 获取子母批作业机台", 4)
    add_cell_paragraph(cell, "通过表fabfutureaction拿取和lot有FutureMerge关系的子批/母批lot，从表r2r_lot_history中（匹配Lotid、productid、layerid）获取最新一笔子批/母批在待判断lot当前layer的作业机台toolid（按实际作业完成时间、记录ID降序取最新记录）。")
    add_heading(cell, "3.2.2.3 判断是否需卡控", 4)
    add_cell_paragraph(cell, "通过r2r_litho_whitelist判断（匹配 productid、layerid、lotid）lot是否为Specify Lot，针对非Specify Lot，若待判断lot的机台与子批/母批作业机台toolid不一致，则卡控Reason：Parent&ChildLotNeedRunSameTool，否则按原逻辑判断。")

    add_heading(cell, "4.LithoAssign新增卡控逻辑", 1)
    add_cell_paragraph(cell, "LithoAssign 增加子母批 Run 相同机台的卡控，并保留原 R2RAutoPirunControl 逻辑。")
    add_heading(cell, "4.1 R2RAutoPirunControl卡控", 2)
    add_cell_paragraph(cell, "针对非Specify Lot，若Pi_SplitFlag='Y'且R2R CD/OVL Status=PIRUNON且Pilot不为Null，则卡控 Reason:R2RAutoPirunControl，否则不卡控。")
    add_heading(cell, "4.2 Parent&ChildLotNeedRunSameTool卡控", 2)
    add_cell_paragraph(cell, "逻辑与 LithoRule 中一致，pretool 获取和子批/母批作业机台获取改为 Central。")

    add_heading(cell, "二、 AMA新增逻辑", 1)
    add_cell_paragraph(cell, "根据 Report:Central_GetLithoR2RAutoPirunInfo 执行物理分批，并将 Pilot 给到 R2R；根据 Report:LithoPiLotAutoDoAdhocSorter执行 TransferFoup。")

    add_heading(cell, "1. 设置Pilot", 2)
    add_cell_paragraph(cell, "获取 Report:Central_GetLithoR2RAutoPirunInfo 栏位信息:Lot、toolid、productid、layerid、reticleid、prereticle、pretool、custom_context_value、pi_splitwafer、IsNeedSplit、isSTNSite。")
    add_heading(cell, "1.1 整批设为Pilot", 3)
    add_cell_paragraph(cell, "当 IsNeedSplit=F 时，直接将整批 Lot 传给 R2R。")

    add_heading(cell, "1.2 物理分批Pilot", 3)
    add_heading(cell, "1.2.1 执行前复核", 4)
    for text in (
        "（1）分批 Lot 属于本厂，即 FAB6 或 FAB8；",
        "（2）Lot 当前状态为 WaitForJobPrep；",
        "（3）Lot 当前站点的 runcardid 为空；",
        "（4）Lot 当前 Capability 为 LithoCapability、L-BARCO-L 或 L-BARCO-S；",
        "（5）CarrierKind='FOUP'；",
        "（6）Report 中选中的 Wafer 确实存在于该 Lot 中。",
    ):
        add_item(cell, text)
    add_cell_paragraph(cell, "仅检查 Wafer 归属关系，不额外检查 Wafer 状态、Slot 或 Chuck。任一复核项不满足时，停止处理该 Lot，记录失败原因并等待下一轮重新计算，不回退为整批 Pilot。")

    add_heading(cell, "1.2.2 空FOUP与MES物理分批", 4)
    add_cell_paragraph(cell, "当 IsNeedSplit=T 时，先执行上述六项复核。复核通过后，先获取并预占空 Foup，再给 MES 物理分批接口，将 pi_splitwafer 从 Lot 中分出。若未拿到可用空 Foup，则不执行物理分批，将整批 Lot 传给 R2R；若分批接口 Fail，则立即释放预占的空 Foup，并将整批 Lot 传给 R2R；否则将分出的子批 pilot 传给 R2R。")
    add_flowchart(cell, "图4  AMA物理分批与回退流程", "04-ama-split.png", width_inches=4.5)

    add_heading(cell, "2. TransferFoup", 2)
    add_cell_paragraph(cell, "获取 Report：LithoPiLotAutoDoAdhocSorter 栏位信息，并按顺序给 MES 打 TransferFOUP 接口，将 Pilot 导到空 FOUP 中，若拿取可用的空 Foup 失败或空 Foup 数量为0，或 MES TransferFOUP 接口失败，则在 AMALog 中记录 Fail 信息。")


def set_metadata_cell(cell, parts):
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_runs(p, parts)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.5

    return doc


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("新增需求申请单")
    set_run_font(run, size=24, italic=True, ascii_font="Times New Roman", east_asia="KaiTi")

    number = doc.add_paragraph()
    number.paragraph_format.space_before = Pt(0)
    number.paragraph_format.space_after = Pt(5)
    number.paragraph_format.line_spacing = 1.0
    run = number.add_run("编号：")
    set_run_font(run)
    run = number.add_run("　　　　　　　　　　　")
    set_run_font(run)
    run.font.underline = True
    run = number.add_run("（此处由信息技术部填写）")
    set_run_font(run)

    table = doc.add_table(rows=10, cols=10)
    table.style = "Table Grid"
    set_table_geometry(table)
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)

    category = table.rows[0].cells[0].merge(table.rows[0].cells[9])
    set_metadata_cell(category, [
        ("类别（请在方框内打勾）：", True, False),
        ("□1. 软件采购　□2. 硬件采购　■3. 功能开发　□4. 工程及服务", False, False),
    ])
    set_cant_split(table.rows[0])

    metadata_cells = []
    for row in table.rows[1:4]:
        left = row.cells[0].merge(row.cells[3])
        right = row.cells[4].merge(row.cells[9])
        metadata_cells.append((left, right))
    set_metadata_cell(metadata_cells[0][0], [("申请部门：", True, False), ("制造部", False, False)])
    set_metadata_cell(metadata_cells[0][1], [("系统名称（类别为3时必填）：", True, False), ("CIM 计算机集成制造系统 Fab6（一科）", False, False)])
    set_metadata_cell(metadata_cells[1][0], [("申请人员：", True, False), ("温浩奇", False, False)])
    set_metadata_cell(metadata_cells[1][1], [("功能模块（类别为3时必填）：", True, False), ("智能派工系统（RTD/DSP）", False, False)])
    set_metadata_cell(metadata_cells[2][0], [("申请日期：", True, False), ("2026-07-24", False, False)])
    set_metadata_cell(metadata_cells[2][1], [("希望交付日期：", True, False), ("2026-07-29", False, False)])
    for row in table.rows[1:4]:
        set_cant_split(row)

    intro = table.rows[4].cells[0].merge(table.rows[4].cells[9])
    clear_cell(intro)
    add_heading(intro, "项目简介和必要性分析：", 1)
    add_cell_paragraph(intro, "当前 LithoAutoSplitPirun 为逻辑分批，分出的 Pilot 与母批在同一 Foup，跨厂场景下会因同 FOUP Lot Transfer 限制，导致 Pilot 无法及时 Pirun，对产线 WIP 流通造成影响。另外，AutoPirun 只针对满足分批条件的 Lot 自动设置为 Pilot 进行 Pirun，不满足条件的 Lot 会被一直卡控，导致产线许多 Lot OverQtime。因此需将 LithoAutoSplitPirun 由逻辑分批修改为物理分批，并优化 Pilot 选择逻辑。")
    set_cant_split(table.rows[4])

    analysis = table.rows[5].cells[0].merge(table.rows[5].cells[9])
    clear_cell(analysis)
    add_heading(analysis, "项目投资方案比较及效果分析：", 1)
    add_heading(analysis, "改善方案：", 2)
    add_cell_paragraph(analysis, "将 LithoAutoSplitPirun 由逻辑分批修改为物理分批，并优化 Pilot 选择逻辑。")
    add_heading(analysis, "效果分析：", 2)
    add_cell_paragraph(analysis, "减少 Lot OverQtime 风险。")
    set_cant_split(table.rows[5])

    content = table.rows[6].cells[0].merge(table.rows[6].cells[9])
    add_content_exact_replica(content)
    set_cant_split(table.rows[6], enabled=False)

    approval_labels = (
        ("申请部门意见：", "申请部门分管领导意见："),
        ("相关部门意见：", "相关部门分管领导意见："),
        ("信息技术部意见：", "信息技术部分管领导意见："),
    )
    for row, labels in zip(table.rows[7:10], approval_labels):
        set_cant_split(row)
        row.height = Mm(27)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        approval_cells = (row.cells[0].merge(row.cells[4]), row.cells[5].merge(row.cells[9]))
        for cell, label in zip(approval_cells, labels):
            clear_cell(cell)
            add_heading(cell, label, 1)
            date_paragraph = add_heading(cell, "日期：", 2)
            date_paragraph.paragraph_format.space_before = Pt(18)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.core_properties.title = "LithoAutoPiRun 新增需求申请单"
    doc.core_properties.subject = "逻辑分批调整为物理分批"
    doc.core_properties.author = "制造部"
    doc.core_properties.keywords = "LithoAutoPiRun, RTD, AMA, Pilot, FOUP, 物理分批"

    doc.save(OUT_PATH)
    shutil.copy2(OUT_PATH, QA_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
