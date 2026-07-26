from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph


def paragraph_image_refs(paragraph) -> list[str]:
    refs: list[str] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rel_id:
            refs.append(rel_id)
    return refs


def paragraph_record(paragraph, relationships) -> dict:
    refs = paragraph_image_refs(paragraph)
    images = []
    for rel_id in refs:
        rel = relationships.get(rel_id)
        if rel is not None:
            images.append(
                {
                    "rel_id": rel_id,
                    "target_ref": str(rel.target_ref),
                    "target_part": str(getattr(rel.target_part, "partname", "")),
                }
            )
    return {
        "style": paragraph.style.name if paragraph.style else "",
        "text": paragraph.text,
        "images": images,
        "xml_has_drawing": bool(paragraph._p.xpath(".//w:drawing")),
        "xml_has_pict": bool(paragraph._p.xpath(".//w:pict")),
    }


def iter_block_items(document):
    para_by_id = {id(p._p): p for p in document.paragraphs}
    table_by_id = {id(t._tbl): t for t in document.tables}
    for child in document.element.body.iterchildren():
        child_id = id(child)
        if child_id in para_by_id:
            yield "paragraph", para_by_id[child_id]
        elif child_id in table_by_id:
            yield "table", table_by_id[child_id]


def inspect(input_path: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)
    relationships = document.part.rels

    blocks: list[dict] = []
    paragraph_index = 0
    table_index = 0
    for block_index, (kind, item) in enumerate(iter_block_items(document)):
        if kind == "paragraph":
            record = paragraph_record(item, relationships)
            blocks.append(
                {
                    "block": block_index,
                    "kind": kind,
                    "paragraph": paragraph_index,
                    **record,
                }
            )
            paragraph_index += 1
        else:
            rows = []
            for row_index, row in enumerate(item.rows):
                rows.append([cell.text for cell in row.cells])
            xml_cells = []
            for row_index, table_row in enumerate(item._tbl.tr_lst):
                for physical_cell_index, table_cell in enumerate(table_row.tc_lst):
                    cell = _Cell(table_cell, item)
                    xml_cells.append(
                        {
                            "row": row_index,
                            "physical_cell": physical_cell_index,
                            "text": cell.text,
                            "paragraphs": [
                                {
                                    "paragraph_in_cell": paragraph_in_cell,
                                    **paragraph_record(paragraph, relationships),
                                }
                                for paragraph_in_cell, paragraph in enumerate(
                                    cell.paragraphs
                                )
                            ],
                            "descendant_paragraphs": [
                                {
                                    "paragraph_in_descendants": paragraph_index,
                                    **paragraph_record(
                                        Paragraph(paragraph_xml, cell),
                                        relationships,
                                    ),
                                }
                                for paragraph_index, paragraph_xml in enumerate(
                                    table_cell.xpath(".//w:p")
                                )
                            ],
                        }
                    )
            blocks.append(
                {
                    "block": block_index,
                    "kind": kind,
                    "table": table_index,
                    "rows": rows,
                    "xml_cells": xml_cells,
                }
            )
            table_index += 1

    report = {
        "input": str(input_path),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": [
            {
                "index": index,
                "width_emu": shape.width,
                "height_emu": shape.height,
                "type": str(shape.type),
            }
            for index, shape in enumerate(document.inline_shapes)
        ],
        "sections": [
            {
                "index": index,
                "page_width_emu": section.page_width,
                "page_height_emu": section.page_height,
                "top_margin_emu": section.top_margin,
                "bottom_margin_emu": section.bottom_margin,
                "left_margin_emu": section.left_margin,
                "right_margin_emu": section.right_margin,
            }
            for index, section in enumerate(document.sections)
        ],
        "blocks": blocks,
    }
    (output_dir / "structure.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    with zipfile.ZipFile(input_path) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                destination = output_dir / "media" / Path(name).name
                destination.parent.mkdir(parents=True, exist_ok=True)
                destination.write_bytes(archive.read(name))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_path", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    inspect(args.input_path, args.output_dir)


if __name__ == "__main__":
    main()
