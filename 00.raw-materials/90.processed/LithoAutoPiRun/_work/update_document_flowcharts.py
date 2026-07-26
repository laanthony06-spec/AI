from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


OLD_CAPTION_PATTERN = re.compile(r"^图[1-4]\s+")


def all_paragraphs(document: Document) -> list[Paragraph]:
    return [
        Paragraph(paragraph_xml, document)
        for paragraph_xml in document.element.body.xpath(".//w:p")
    ]


def image_relationship_ids(paragraph: Paragraph) -> list[str]:
    relationship_ids: list[str] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        relationship_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if relationship_id:
            relationship_ids.append(relationship_id)
    return relationship_ids


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def insert_picture_after(
    document: Document,
    target_prefix: str,
    image_path: Path,
    width_inches: float,
    alt_text: str,
) -> None:
    target = next(
        (
            paragraph
            for paragraph in all_paragraphs(document)
            if paragraph.text.strip().startswith(target_prefix)
        ),
        None,
    )
    if target is None:
        raise RuntimeError(f"未找到插图位置：{target_prefix}")

    paragraph_xml = OxmlElement("w:p")
    target._p.addnext(paragraph_xml)
    picture_paragraph = Paragraph(paragraph_xml, document)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(5)
    picture_paragraph.paragraph_format.space_after = Pt(5)
    picture_paragraph.paragraph_format.keep_together = True

    run = picture_paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width_inches))
    inline_shape._inline.docPr.set("name", image_path.stem)
    inline_shape._inline.docPr.set("descr", alt_text)


def update_document(
    source_path: Path,
    output_path: Path,
    pilot_image: Path,
    wait_image: Path,
    parent_child_image: Path,
) -> None:
    document = Document(source_path)
    paragraphs = all_paragraphs(document)

    remove_set: set[int] = set()
    removed_relationships: set[str] = set()
    for index, paragraph in enumerate(paragraphs):
        if paragraph._p.xpath(".//w:drawing | .//w:pict"):
            remove_set.add(index)
            removed_relationships.update(image_relationship_ids(paragraph))
            if index + 1 < len(paragraphs) and not paragraphs[index + 1].text.strip():
                remove_set.add(index + 1)
        elif OLD_CAPTION_PATTERN.match(paragraph.text.strip()):
            remove_set.add(index)

    for index in sorted(remove_set, reverse=True):
        remove_paragraph(paragraphs[index])

    for relationship_id in removed_relationships:
        if relationship_id in document.part.rels:
            document.part.drop_rel(relationship_id)

    insert_picture_after(
        document,
        "将排序第一的 Lot+Context 固定",
        pilot_image,
        6.2,
        "1.4.3 Context 循环挑选 Pilot 逻辑流程图",
    )
    insert_picture_after(
        document,
        "（2）若 LithoPilot 不在 AdhocSoter 站点",
        wait_image,
        6.8,
        "3.1 WaitPilotChangeFOUP 卡控与解除逻辑流程图",
    )
    insert_picture_after(
        document,
        "通过r2r_litho_whitelist判断",
        parent_child_image,
        6.5,
        "3.2.2 Parent&ChildLotNeedRunSameTool 卡控逻辑流程图",
    )

    document.core_properties.last_modified_by = "Codex"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_path", type=Path)
    parser.add_argument("output_path", type=Path)
    parser.add_argument("pilot_image", type=Path)
    parser.add_argument("wait_image", type=Path)
    parser.add_argument("parent_child_image", type=Path)
    args = parser.parse_args()
    update_document(
        args.source_path,
        args.output_path,
        args.pilot_image,
        args.wait_image,
        args.parent_child_image,
    )


if __name__ == "__main__":
    main()
