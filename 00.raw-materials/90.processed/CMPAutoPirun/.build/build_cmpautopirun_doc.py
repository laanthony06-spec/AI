from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Obsidian\work\OBSidianCodex")
OUT_DIR = ROOT / "00.raw-materials" / "90.processed" / "CMPAutoPirun"
BUILD_DIR = OUT_DIR / ".build"
OUT = OUT_DIR / "CMPAutoPirun_详细设计说明书_V1.0.docx"

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "1F6D7A"
GOLD = "B8860B"
INK = "1F2937"
GRAY = "5B6573"
LIGHT = "E8EEF5"
LIGHTER = "F4F6F9"
BORDER = "C9D2DE"
WHITE = "FFFFFF"
RISK = "9B1C1C"
OK = "1E6B45"

BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"


def rgb(hex_string):
    return RGBColor.from_string(hex_string)


def set_run_font(run, name=BODY_FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths_dxa):
                continue
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)


def style_table(table, header=True, first_col_bold=False, font_size=9.2):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    border_args = {"val": "single", "sz": "6", "color": BORDER, "space": "0"}
    for r_idx, row in enumerate(table.rows):
        set_row_cant_split(row)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, top=border_args, start=border_args, bottom=border_args, end=border_args)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT)
            else:
                set_cell_shading(cell, WHITE)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.18
                for run in p.runs:
                    set_run_font(run, size=font_size, color=INK,
                                 bold=(header and r_idx == 0) or (first_col_bold and c_idx == 0))
        if r_idx == 0 and header:
            set_repeat_table_header(row)


def set_paragraph_border(paragraph, edge="bottom", color=BLUE, size=12, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    node = OxmlElement(f"w:{edge}")
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), color)
    p_bdr.append(node)


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]

    def add_scheme(fmt, text_value, bullet=False):
        abstract_id = max(abstract_ids or [0]) + 1
        abstract_ids.append(abstract_id)
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
        if bullet:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Symbol")
            fonts.set(qn("w:hAnsi"), "Symbol")
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num_id = max(num_ids or [0]) + 1
        num_ids.append(num_id)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)
        return num_id

    return add_scheme("bullet", "", True), add_scheme("decimal", "%1.")


def clone_num_instance(doc, source_num_id):
    numbering = doc.part.numbering_part.element
    source = None
    num_ids = []
    for num in numbering.findall(qn("w:num")):
        num_id = int(num.get(qn("w:numId")))
        num_ids.append(num_id)
        if num_id == source_num_id:
            source = num
    if source is None:
        raise ValueError(f"numbering instance {source_num_id} not found")
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max(num_ids or [0]) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_id))
    abs_node = OxmlElement("w:abstractNumId")
    abs_node.set(qn("w:val"), abstract_id)
    new_num.append(abs_node)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)
    return new_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, italic=italic)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_row_cant_split(table.rows[0])
    set_cell_shading(cell, "F7F8FA")
    set_cell_border(cell, top={"val": "single", "sz": "8", "color": BORDER},
                    start={"val": "single", "sz": "20", "color": TEAL},
                    bottom={"val": "single", "sz": "8", "color": BORDER},
                    end={"val": "single", "sz": "8", "color": BORDER})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, name=CODE_FONT, size=8.8, color="263238")
        if idx != len(lines) - 1:
            r.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, title, text, color=TEAL, fill="EEF6F7"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_row_cant_split(table.rows[0])
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={"val": "nil"}, start={"val": "single", "sz": "24", "color": color},
                    bottom={"val": "nil"}, end={"val": "nil"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title + "  ")
    set_run_font(r1, size=10.2, color=color, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9.1, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    set_table_geometry(table, widths)
    style_table(table, header=True, first_col_bold=first_col_bold, font_size=font_size)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return table


def _diagram_font(size, bold=False):
    choices = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in choices:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _center_text(draw, xy, text, font, fill, spacing=5):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill,
                        align="center", spacing=spacing)


def _arrow(draw, start, end, fill="#667085", width=4):
    draw.line([start, end], fill=fill, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 14
    spread = .55
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=fill)


def draw_architecture(path):
    img = Image.new("RGB", (1836, 810), "white")
    draw = ImageDraw.Draw(img)
    title_font = _diagram_font(38, True)
    text_font = _diagram_font(25, True)
    _center_text(draw, (918, 70), "CMPAutoPirun 逻辑架构", title_font, "#" + NAVY)
    boxes = [
        (50, 300, 285, 500, "调度触发\nAMA WatchDog", NAVY),
        (385, 150, 740, 655, "输入数据\nCMPAssign\nR2R 配置\nRTDConfig\n机台计数器", BLUE),
        (805, 280, 1115, 525, "候选选择器\nLot / Pilot / STN", TEAL),
        (1190, 280, 1440, 525, "分批执行\nSplit", GOLD),
        (1510, 170, 1790, 350, "R2R 接口\nPi-Run 信息", NAVY),
        (1510, 480, 1790, 660, "日志与监控\n审计链路", GRAY),
    ]
    for x1, y1, x2, y2, text_value, color in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="white", outline="#" + color, width=5)
        _center_text(draw, ((x1+x2)/2, (y1+y2)/2), text_value, text_font, "#" + color, spacing=7)
    for a, b in [((335,400),(385,400)), ((740,400),(805,400)), ((1115,400),(1190,400)),
                 ((1440,360),(1510,260)), ((1440,455),(1510,570))]:
        _arrow(draw, a, b)
    img.save(path, quality=95)


def draw_flow(path):
    img = Image.new("RGB", (1836, 1152), "white")
    draw = ImageDraw.Draw(img)
    box_font = _diagram_font(24, True)
    small_font = _diagram_font(21, True)
    label_font = _diagram_font(20, True)

    def box(cx, cy, w, h, text, color=NAVY):
        x1, y1, x2, y2 = cx-w/2, cy-h/2, cx+w/2, cy+h/2
        draw.rounded_rectangle((x1,y1,x2,y2), radius=16, fill="white", outline="#"+color, width=4)
        _center_text(draw, (cx,cy), text, box_font, "#"+color, spacing=5)

    def diamond(cx, cy, w, h, text, color=GOLD):
        pts = [(cx,cy-h/2),(cx+w/2,cy),(cx,cy+h/2),(cx-w/2,cy)]
        draw.polygon(pts, fill="#FFF9E8", outline="#"+color)
        draw.line(pts+[pts[0]], fill="#"+color, width=4)
        _center_text(draw, (cx,cy), text, small_font, "#5F4600", spacing=4)

    box(180,120,250,100,"WatchDog\n按配置触发")
    box(540,120,300,100,"读取配置与\n最新 CMPAssign")
    diamond(940,120,275,150,"配置有效且\n存在候选 Lot？")
    box(1395,120,330,100,"按 selfcapability\n分组")
    diamond(1395,350,350,165,"Q-time / Lifetime\n风险？")
    box(950,350,320,110,"needpirunflag = 2\n风险优先")
    diamond(1395,600,350,165,"Highwip 条件\n满足？")
    box(950,600,320,110,"needpirunflag = 1\n宽度需求")
    box(1395,840,350,112,"选择可 Pi 机台\n低 loading 优先")
    box(950,840,320,112,"按配置片数\n执行子批 Split")
    diamond(560,840,280,150,"Split 成功？")
    box(190,840,260,108,"发送 R2R\n记录审计日志",TEAL)
    box(560,1050,290,90,"失败补偿 / 告警",RISK)
    box(950,1050,280,90,"本轮结束",GRAY)

    arrows = [
        ((305,120),(390,120),None), ((690,120),(802,120),None), ((1077,120),(1230,120),"是"),
        ((1395,170),(1395,267),None), ((1220,350),(1110,350),"是"), ((1395,433),(1395,517),"否"),
        ((1220,600),(1110,600),"是"), ((1395,683),(1395,784),"是"), ((1220,840),(1110,840),None),
        ((790,840),(700,840),None), ((420,840),(320,840),"是"), ((560,915),(560,1005),"否"),
        ((815,350),(950,785),None), ((815,600),(950,785),None), ((1075,120),(950,1005),"否 / 无候选"),
        ((320,895),(815,1035),"成功后继续下一组"),
    ]
    for a,b,label in arrows:
        _arrow(draw,a,b)
        if label:
            mx,my=(a[0]+b[0])/2,(a[1]+b[1])/2
            bbox=draw.textbbox((0,0),label,font=label_font)
            pad=5
            draw.rectangle((mx-(bbox[2]-bbox[0])/2-pad,my-(bbox[3]-bbox[1])/2-pad,
                            mx+(bbox[2]-bbox[0])/2+pad,my+(bbox[3]-bbox[1])/2+pad),fill="white")
            draw.text((mx-(bbox[2]-bbox[0])/2,my-(bbox[3]-bbox[1])/2),label,font=label_font,fill="#667085")
    img.save(path, quality=95)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(6.35))
    p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=GRAY, italic=True)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(23)
    section.right_margin = Mm(22)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (11.5, NAVY, 10, 5),
    }
    for level, (size, color, before, after) in heading_specs.items():
        style = styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    lp = styles["List Paragraph"]
    lp.font.name = BODY_FONT
    lp._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    lp.font.size = Pt(10.5)
    lp.paragraph_format.left_indent = Inches(.375)
    lp.paragraph_format.first_line_indent = Inches(-.188)
    lp.paragraph_format.space_after = Pt(4)
    lp.paragraph_format.line_spacing = 1.25

    if "TOC Heading" in styles:
        toc_h = styles["TOC Heading"]
        toc_h.font.name = BODY_FONT
        toc_h._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r1 = hp.add_run("CMPAutoPirun 详细设计说明书")
    set_run_font(r1, size=8.5, color=GRAY, bold=True)
    r2 = hp.add_run("  |  Confidential II")
    set_run_font(r2, size=8.5, color=GRAY)

    footer = section.footer
    ft = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ft, [6200, 3160], indent_dxa=0)
    ft.cell(0,0).text = "AMA-DDS-CMP-AP-001  ·  V1.0"
    left = ft.cell(0,0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in left.runs:
        set_run_font(run, size=8.2, color=GRAY)
    right = ft.cell(0,1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("第 ")
    set_run_font(r, size=8.2, color=GRAY)
    set_run_font(add_field(right, "PAGE", "1"), size=8.2, color=GRAY)
    r = right.add_run(" 页 / 共 ")
    set_run_font(r, size=8.2, color=GRAY)
    set_run_font(add_field(right, "NUMPAGES", "1"), size=8.2, color=GRAY)
    r = right.add_run(" 页")
    set_run_font(r, size=8.2, color=GRAY)
    for cell in ft.rows[0].cells:
        set_cell_border(cell, top={"val":"single","sz":"8","color":BORDER},
                        start={"val":"nil"}, bottom={"val":"nil"}, end={"val":"nil"})
        set_cell_margins(cell, top=80, bottom=0, start=0, end=0)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AMA · CMP R2R")
    set_run_font(r, size=12, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CMPAutoPirun")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("CMP R2R Pi-Run Lot 自动分批")
    set_run_font(r, size=17, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    set_paragraph_border(p, edge="top", color=TEAL, size=18, space=10)
    r = p.add_run("详细设计说明书")
    set_run_font(r, size=15, color=INK, bold=True)
    set_paragraph_border(p, edge="bottom", color=TEAL, size=18, space=10)

    meta = [
        ("文档编号", "AMA-DDS-CMP-AP-001"),
        ("版本", "V1.0"),
        ("状态", "评审稿"),
        ("编制日期", "2026-07-10"),
        ("保密级别", "Confidential II / 商密二级"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    for i, (k, v) in enumerate(meta):
        table.cell(i,0).text = k
        table.cell(i,1).text = v
    set_table_geometry(table, [2400, 5000], indent_dxa=980)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            set_cell_border(cell, top={"val":"nil"}, start={"val":"nil"},
                            bottom={"val":"single","sz":"6","color":BORDER}, end={"val":"nil"})
            set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=9.8, color=GRAY if j == 0 else INK, bold=(j == 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("依据需求单、测试用例与现有技术文档截图整理")
    set_run_font(r, size=9, color=GRAY, italic=True)
    doc.add_page_break()


def build_document():
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    arch = BUILD_DIR / "architecture.png"
    flow = BUILD_DIR / "main_flow.png"
    draw_architecture(arch)
    draw_flow(flow)

    doc = Document()
    configure_styles(doc)
    setup_header_footer(doc)
    bullet_id, decimal_id = create_numbering(doc)
    add_cover(doc)

    add_heading(doc, "版本历史", 1)
    add_table(doc,
              ["版本", "日期", "状态", "变更说明", "编制"],
              [["V1.0", "2026-07-10", "评审稿", "基于技术文档、需求单、PPT 与测试用例截图形成首版详细设计。", "项目组"]],
              [900, 1500, 1100, 4360, 1500], font_size=9.2)
    add_heading(doc, "文档控制", 2)
    add_table(doc,
              ["项目", "说明"],
              [
                  ["文档用途", "用于 CMPAutoPirun 功能评审、开发实现、测试验收与运维交接。"],
                  ["适用对象", "AMA / R2R 开发、CMP 工艺、测试、运维及系统接口负责人。"],
                  ["信息来源", "技术文档1-4.jpg、需求单1-5.jpg、Testcase1-2.jpg、PPT.jpg。"],
                  ["保留原则", "截图中可辨识的参数名、表名与业务术语按原文保留；无法确认的实现细节统一列入“待确认事项”。"],
              ], [2100, 7260], font_size=9.4, first_col_bold=True)
    add_callout(doc, "阅读提示", "本文将截图中的分散规则统一为可执行的决策链。带“待确认”的内容不应直接作为生产配置，应在设计评审后关闭。")
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("目录")
    set_run_font(r, size=20, color=NAVY, bold=True)
    toc_entries = [
        ("版本历史", "2", 0),
        ("文档说明", "4", 0),
        ("业务需求", "4", 0),
        ("总体设计", "5", 0),
        ("配置设计", "6", 0),
        ("数据与接口设计", "8", 0),
        ("核心算法设计", "9", 0),
        ("调度、事务与异常处理", "11", 0),
        ("日志、监控与安全", "12", 0),
        ("测试与验收", "13", 0),
        ("上线与回退", "14", 0),
        ("待确认事项", "14", 1),
        ("附录：伪代码与来源清单", "15", 0),
    ]
    for title, page, level in toc_entries:
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.left_indent = Inches(0.28 * level)
        toc_p.paragraph_format.space_before = Pt(0)
        toc_p.paragraph_format.space_after = Pt(5 if level == 0 else 3)
        toc_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = toc_p.add_run(f"{title}\t{page}")
        set_run_font(run, size=10.5 if level == 0 else 9.8, color=INK, bold=(level == 0))
    doc.add_page_break()

    add_heading(doc, "文档说明", 1)
    add_heading(doc, "目的", 2)
    add_body(doc, "本文定义 CMPAutoPirun 的业务目标、配置入口、数据依赖、Lot/Pilot/STN 选择算法、机台筛选、子批 Split、R2R 通知、异常补偿及测试验收要求，为实现与评审提供统一基线。")
    add_heading(doc, "范围", 2)
    for item in [
        "包含：AMA WatchDog 定时触发、CMP Pi-Run Lot 自动选择、机台分配、物理分批、R2R 信息发送与全链路日志。",
        "包含：Lowwip 兼容逻辑与 Highwip 新增逻辑，以及 Q-time / Lifetime 风险优先策略。",
        "不包含：CMPAssign 原有调度算法重构、R2R 内部派工策略、设备控制器本体改造。",
    ]:
        add_list_item(doc, item, bullet_id)
    add_heading(doc, "术语与缩写", 2)
    add_table(doc, ["术语", "含义"], [
        ["Pi-Run / Pirun", "CMP 机台在正式生产前执行的先导/试跑作业。"],
        ["Lot / 子批", "生产批次及由物理 Split 产生的子批。"],
        ["Pilot", "被选为 Pi-Run 的 Lot 或其分批子批。"],
        ["STN / Side", "Lot 可运行的机台边或站点标识。"],
        ["selfcapability", "机台能力分组，用于 Lot 与可运行机台匹配。"],
        ["Q-time", "工序时间窗约束；超时或接近超时会产生风险。"],
        ["Lifetime", "Pad/耗材等寿命计数约束。"],
        ["needpirunflag", "Pi-Run 需求标识：0=不需要，1=宽度/高 WIP 需求，2=Q-time 或 Lifetime 风险。"],
    ], [2000, 7360], font_size=9.2, first_col_bold=True)

    add_heading(doc, "业务需求", 1)
    add_heading(doc, "背景", 2)
    add_body(doc, "CMP 堆货时，PE 手工执行 Pi-Run 效率较低，容易造成 Lot 超过 Q-time。需要由 AMA 根据 Lot 作业条件、机台 loading 与耗材 Lifetime 状态自动选择 Pi-Run Lot 并分批，从而降低人工操作和等待时间。")
    add_heading(doc, "业务目标与收益", 2)
    for item in [
        "自动发现需要 Pi-Run 的风险 Lot 与宽度需求 Lot。",
        "将 Lot 分配到 loading 较低且 Lifetime 充足的可运行机台。",
        "自动生成符合配置片数的子批，并把 Pi-Run 信息发送至 R2R。",
        "减少 over Q-time Lot 数量；参考项目材料，目标为机台生产效率提升 3% 以上。",
        "保证每次选择与分批均可追溯、可重放、可审计。",
    ]:
        add_list_item(doc, item, bullet_id)
    add_heading(doc, "关键业务约束", 2)
    add_table(doc, ["约束", "设计要求"], [
        ["唯一性", "同一轮中每个 Lot 仅可被选择一次；每台机台同一时刻仅分配一个 Pi-Run 子批。"],
        ["冷却期", "Lot 被选中后 1 天内不再参与后续选择；STN 在 pilot 状态解除前不再参与循环。"],
        ["实时性", "候选 Lot 必须基于最新 sysId 结果，并在决策前与实时记录进行 double check。"],
        ["兼容性", "ControlType=Lowwip 沿用原逻辑；ControlType=Highwip 启用新增 Pi-Run 选择逻辑。"],
        ["安全失败", "关键配置缺失、数据不一致或 Split/R2R 调用失败时，不得继续产生新的不可追溯子批。"],
    ], [1800, 7560], font_size=9.2, first_col_bold=True)

    add_heading(doc, "总体设计", 1)
    add_heading(doc, "逻辑架构", 2)
    add_figure(doc, arch, "图 1  CMPAutoPirun 逻辑架构")
    add_body(doc, "AMA WatchDog 负责调度；候选选择器聚合 CMPAssign、R2R 配置、RTDConfig 与机台 Lifetime/Loading 数据；分批执行器完成物理 Split；结果同步给 R2R，并写入审计日志与监控指标。")
    add_heading(doc, "主流程", 2)
    add_figure(doc, flow, "图 2  CMPAutoPirun 主决策流程")
    add_callout(doc, "优先级原则", "needpirunflag=2 的 Q-time / Lifetime 风险 Lot 高于 needpirunflag=1 的 Highwip 宽度需求 Lot；同优先级内再按 Remaining、loading 与 Lifetime 频度排序。", color=GOLD, fill="FFF9E8")

    add_heading(doc, "配置设计", 1)
    add_heading(doc, "AMA.TriggerConfig", 2)
    add_table(doc, ["配置项", "建议值 / 规则", "校验"], [
        ["FunctionName", "CMPAutoSplitPirunLot", "精确匹配；未配置时不执行。"],
        ["Switch", "Y", "仅 Y 启用；其他值按关闭处理。"],
        ["Trigger Time Slot", "00:00-23:59", "当前时间必须落在范围内。"],
        ["Trigger Count/Time", "5（分钟）", "WatchDog 每 5 分钟触发一次。"],
    ], [2200, 3200, 3960], font_size=9.2)
    add_code_block(doc, [
        "FunctionName = CMPAutoSplitPirunLot",
        "Switch = Y",
        "TriggerTimeSlot = 00:00-23:59",
        "TriggerIntervalMinutes = 5",
    ])
    add_heading(doc, "RTD.CMPlowWIP", 2)
    add_table(doc, ["字段", "规则", "用途"], [
        ["Capability", "必填，不支持模糊匹配", "能力范围匹配。"],
        ["SelfCapability", "必填，不支持模糊匹配", "Pi-Run 分组与机台边匹配。"],
        ["Ratio", "必填，不支持模糊匹配", "Highwip LimitWIP 计算。"],
        ["ControlType", "必填；Lowwip / Highwip；不支持模糊匹配", "Lowwip 兼容原逻辑，Highwip 启用新增逻辑。"],
    ], [1700, 4260, 3400], font_size=9.1)
    add_callout(doc, "兼容策略", "历史配置若未显式设置 ControlType，按原需求描述兼容为 Lowwip；正式上线前建议完成配置数据清洗并改为显式值。")
    add_heading(doc, "R2R 与 Lifetime 配置", 2)
    add_table(doc, ["配置/数据", "用途", "要求"], [
        ["R2R tech / prod / stage / count / qty", "确定分批数量与工艺匹配", "多条 Count 时取最小；按 selfcapability 的 qty 处理。"],
        ["mergestep / futuremerge", "子批回并控制", "与所选 Pi-Run Lot 的工艺路径一致。"],
        ["R2R_IAPC", "分批片数", "仅当 Lot qty 大于配置片数时允许 Split。"],
        ["CMPLowWipMachinedouble", "特定边 Lifetime 特判", "参数名大小写与多值分隔符需在评审中确认。"],
    ], [2300, 3200, 3860], font_size=9.1)

    add_heading(doc, "数据与接口设计", 1)
    add_heading(doc, "输入数据", 2)
    add_table(doc, ["来源", "关键字段/数据", "使用方式"], [
        ["CMPAssign 最新结果", "sysId、Lot 状态、qty、QTimeWorse、LifeTimeWorse、Remaining、selfcapability、assignedstn、reasoncode、totalwip", "筛选候选 Lot、实时复核、计算优先级与 loading。"],
        ["tb_machineselfcapa", "tool_name（按 side）、selfcapability", "取得机台边与 selfcapability 的对应关系。"],
        ["R2R 配置表", "tech、prod、stage、count、qty、mergestep、futuremerge", "计算 Pi-Run 上限、Split 数量与回并信息。"],
        ["MFGCIM.tb_selfcapa_rule", "tech/prod/stage 到 selfcapability 映射", "补充 CMPAssign 未返回的 selfcapability。"],
        ["RTD.CMPlowWIP", "Capability、SelfCapability、Ratio、ControlType", "确定 Lowwip/Highwip 策略与 LimitWIP。"],
        ["设备计数数据", "equipmentid、chamberid、metertype、current(value)、max_value", "评估 Lifetime 风险。"],
    ], [2100, 3860, 3400], font_size=8.85)
    add_heading(doc, "建议输出接口契约（待双方确认）", 2)
    add_table(doc, ["字段", "说明", "必填"], [
        ["runId / traceId", "本次 WatchDog 运行与链路追踪标识", "是"],
        ["parentLotId / childLotId", "母批与 Pi-Run 子批标识", "是"],
        ["needPirunFlag", "0 / 1 / 2 决策结果", "是"],
        ["selfCapability", "能力分组", "是"],
        ["targetStn / targetTool", "目标机台边与机台", "是"],
        ["splitQty", "子批片数", "是"],
        ["mergeStep / futureMerge", "回并配置", "是"],
        ["reasonCode", "选择原因及 PE Pi-Run 原因", "是"],
        ["decisionTime", "决策时间", "是"],
    ], [2300, 5560, 1500], font_size=9.0)
    add_callout(doc, "接口原则", "接口必须支持幂等键（建议 runId + parentLotId + targetStn）与明确的成功/失败码；R2R 确认前不得把本地状态标记为最终完成。")

    add_heading(doc, "核心算法设计", 1)
    add_heading(doc, "候选 Lot 初筛与实时复核", 2)
    for item in [
        "按 sysId 读取最新 CMPAssign 结果，并取得 Lot 当前状态、数量、Q-time/Lifetime 风险、Remaining 与可运行 STN。",
        "仅保留状态可用且 qty > 5 的 Lot；具体可用状态白名单由 CMPAssign 接口方确认。",
        "在进入决策前对 CMPLotAssignment 记录进行 double check；若记录已变化，放弃缓存并使用实时数据重新计算。",
        "按照 Lot 的 tech/prod/stage 与 R2R 配置匹配 selfcapability、mergestep、qty 与 count；多条 Count 配置取最小值。",
    ]:
        add_list_item(doc, item, decimal_id)

    add_heading(doc, "Pi-Run 需求判定", 2)
    add_table(doc, ["优先级", "触发条件", "needpirunflag", "选择规则"], [
        ["P0", "QTimeWorse=T 或 LifeTimeWorse=T", "2", "风险 Lot 优先；同 selfcapability 内按风险类型与 Remaining 排序。"],
        ["P1", "ControlType=Highwip 且 QWIP > LimitWIP，且原风险逻辑未命中", "1", "从对应 selfcapability 的高 loading 场景选择适合 Pi-Run 的 Lot。"],
        ["P2", "以上均不满足", "0", "不参与本轮 Pi-Run。"],
    ], [1000, 3500, 1500, 3360], font_size=8.8)
    add_code_block(doc, [
        "LimitWIP = SumSideWIP × Min(QTimeLimit) × Ratio",
        "if QTimeWorse or LifeTimeWorse: needpirunflag = 2",
        "elif ControlType == 'Highwip' and QWIP > LimitWIP: needpirunflag = 1",
        "else: needpirunflag = 0",
    ])
    add_callout(doc, "待确认", "截图中 Remaining 的空值处理、风险 Lot 的精确排序方向，以及 LimitWIP 公式中各字段单位需由业务与开发共同确认。", color=RISK, fill="FFF1F1")

    add_heading(doc, "Pilot 与机台选择", 2)
    for item in [
        "根据 Lot 的可运行 STN 列表建立候选机台集合，并使用 Reasoncode_CMP 过滤。截图可辨识的 PE Pi-Run reason 包括 APC_JOB_OFF、APC_JOB_OFF;AssignOtherSTN、APC_JOB_OFF;AssignOtherSide。",
        "过滤机况不符合要求、存在冲突 pilot、预计执行后会超过 Lifetime 频度的机台边。",
        "在剩余机台中选择 loading 最低者；loading 相同时选择 Lifetime 频度较低者。",
        "若目标机台已存在 pilot，重新执行 pilot/机台选择；若无满足条件机台，本组安全结束并记录原因。",
        "每台机台同一时刻仅允许一个 Pi-Run 子批；已选 STN 在 pilot 解除前不再参与循环。",
    ]:
        add_list_item(doc, item, decimal_id)
    add_heading(doc, "Lifetime 风险计算", 2)
    add_table(doc, ["步骤", "规则"], [
        ["计数器读取", "按 equipmentid、chamberid、metertype 读取 mesprod.fabcmrpwafercount、mesprod.freqequipment、mesprod.xsiteeqpusagemeter。"],
        ["计数类型", "Table Polishing Wafer Count、PadReplaceWfrCnt、Pad wafer count。"],
        ["边级 Current", "同一机台边包含多个 Chamber 时，取对应类型 Current(value) 的最大值。"],
        ["边级阈值", "取各 Chamber max_value 的最小值，形成保守阈值 Min(max_value)。"],
        ["预测值", "New(lifetime) = 当前 Current(value) + 待运行 Lot Qty。"],
        ["风险判定", "若 New(lifetime) > Min(max_value)，则该机台边存在超 Lifetime 风险并被剔除。"],
        ["特殊边", "对于 C-CCCU-C.T 的部分边，若命中 RTDConfig 参数 CMPLowWipMachinedouble，则使用 Qty/2 参与预测。"],
    ], [1800, 7560], font_size=8.95, first_col_bold=True)
    add_heading(doc, "子批 Split 与循环终止", 2)
    for item in [
        "按 R2R_IAPC 配置片数生成 Pi-Run 子批；母批 qty 必须大于 Split 片数。",
        "为母批设置 futuremerge，并记录 mergestep；该 Lot 与 STN 不再参加本轮后续选择。",
        "下一次 Assign 对拆分子批进行卡控：其他可作业 STN 需以 R2RPirunlot 原因限制，Pi-Run 子批在规则中移除 LowWipControl。",
        "按 selfcapability 分组循环选择，直到达到该组配置上限，或已无可选 Lot/STN。",
        "每个 Lot 每轮仅可被选择一次；冷却期内不得重复进入候选集合。",
    ]:
        add_list_item(doc, item, decimal_id)

    add_heading(doc, "调度、事务与异常处理", 1)
    add_heading(doc, "WatchDog 调度", 2)
    add_body(doc, "当 TriggerConfig 配置有效且 Switch=Y 时，WatchDog 在时间窗内每 5 分钟执行一次。每次运行必须生成唯一 runId，并在执行前获取防重入锁；上一次任务未结束时，本次触发应跳过或排队，不能并发操作同一 Lot。")
    add_heading(doc, "事务边界", 2)
    add_table(doc, ["阶段", "提交点", "失败处理"], [
        ["决策", "候选与目标机台确定后写入待执行记录", "不产生物理变化；释放锁并记录原因。"],
        ["Split", "母批/子批关系创建成功", "部分成功时执行补偿或转人工，不得重复 Split。"],
        ["R2R 通知", "收到 R2R 明确成功回执", "按幂等键重试；超过阈值告警并保留待补发状态。"],
        ["完成", "本地状态、R2R 状态与审计记录一致", "对账任务识别不一致并触发修复。"],
    ], [1600, 3760, 4000], font_size=9.0)
    add_heading(doc, "异常处理矩阵", 2)
    add_table(doc, ["异常", "系统行为", "监控/告警"], [
        ["关键配置缺失或非法", "终止当前组，不做 Split", "配置错误告警，带配置键与 selfcapability。"],
        ["CMPAssign 数据已变化", "重新读取并重算", "记录 stale_data 次数。"],
        ["无候选 Lot / STN", "正常结束当前组", "仅记录业务指标，不告警。"],
        ["并发锁冲突", "跳过本次运行", "记录 skipped_due_to_lock。"],
        ["Split 失败", "停止后续步骤，执行补偿", "高优先级告警，附 parentLotId/runId。"],
        ["R2R 超时或失败", "幂等重试，转待补发", "接口告警与重试次数。"],
        ["日志写入失败", "不影响已完成的设备动作，但必须落本地补偿队列", "审计完整性告警。"],
    ], [1900, 3860, 3600], font_size=8.85)

    add_heading(doc, "日志、监控与安全", 1)
    add_heading(doc, "审计日志", 2)
    add_body(doc, "每次选择均需记录“输入快照—过滤原因—排序分值—最终选择—Split 结果—R2R 回执”的完整决策链。敏感业务数据按商密二级要求控制访问，不在普通应用日志中输出完整生产配方或无关个人信息。")
    add_table(doc, ["日志域", "最少记录项"], [
        ["运行级", "runId、触发时间、配置版本、开始/结束时间、状态、耗时。"],
        ["Lot 级", "LotId、sysId、qty、selfcapability、风险标识、Remaining、过滤原因。"],
        ["机台级", "STN/tool、loading、Lifetime current/max、reasoncode、是否已有 pilot。"],
        ["执行级", "parent/child Lot、splitQty、mergeStep、futureMerge、幂等键、返回码。"],
    ], [1900, 7460], font_size=9.1, first_col_bold=True)
    add_heading(doc, "关键指标", 2)
    for item in [
        "WatchDog 成功率、平均/最大执行时长、重入跳过次数。",
        "每轮候选 Lot 数、flag=1/2 数、成功 Split 数、无可用机台数。",
        "Split 成功率、R2R 接口成功率与重试率、对账差异数。",
        "over Q-time Lot 数量变化、Pi-Run 等待时长、机台生产效率提升比例。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "测试与验收", 1)
    add_heading(doc, "核心测试场景", 2)
    test_rows = [
        ["T01", "selfcapability 分组", "按 tool_name/side 取得 selfcapability，并按组取 picountconfig 最小值。", "分组与上限正确。"],
        ["T02", "候选初筛", "Lot 状态可用且 qty>5。", "仅保留符合条件的 Lot。"],
        ["T03", "实时复核", "缓存记录与 CMPLotAssignment 实时数据不一致。", "使用实时数据重算。"],
        ["T04", "风险优先", "QTimeWorse 或 LifeTimeWorse 为 T。", "needpirunflag=2，优先于 flag=1。"],
        ["T05", "Highwip", "QWIP>LimitWIP 且无原风险。", "needpirunflag=1。"],
        ["T06", "机台过滤", "机况异常/已有 pilot/预计超 Lifetime。", "不可 Pi 的机台被剔除。"],
        ["T07", "机台排序", "多个机台可用。", "选 loading 最低；相同则选 Lifetime 频度低。"],
        ["T08", "唯一性与冷却", "Lot/STN 已在本轮选中或仍在冷却期。", "不重复选择。"],
        ["T09", "分批片数", "母批 qty 大于 R2R_IAPC 配置片数。", "正确 Split 并输出 componentinfo。"],
        ["T10", "调度开关", "Switch=N 或不在时间窗。", "不执行；不产生子批。"],
        ["T11", "R2R 失败", "接口超时/返回失败。", "幂等重试，最终进入待补发并告警。"],
        ["T12", "配置兼容", "ControlType 为空、Lowwip、Highwip、非法值。", "兼容/启用/拒绝逻辑符合设计。"],
    ]
    add_table(doc, ["ID", "场景", "输入/前置条件", "预期结果"], test_rows,
              [800, 1600, 4000, 2960], font_size=8.45)
    add_heading(doc, "验收标准", 2)
    for item in [
        "功能：截图需求与测试用例中可辨识的选择、机台分配、Split 和 R2R 发送逻辑全部通过。",
        "一致性：同一 Lot、同一机台不存在重复分批；异常恢复后数据可对账。",
        "性能：单轮执行在约定窗口内完成，且不与下一轮 WatchDog 重叠。",
        "效果：over Q-time Lot 数量下降；机台生产效率提升目标按上线前基线与观察周期进行验证，参考目标为 3% 以上。",
        "可运维：所有决策可通过 runId 追溯，关键错误具备告警、重试与人工处置入口。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "上线与回退", 1)
    add_heading(doc, "上线步骤", 2)
    rollout_decimal_id = clone_num_instance(doc, decimal_id)
    for item in [
        "完成配置清洗与接口字段确认，关闭全部待确认事项。",
        "在影子模式运行，仅输出决策不执行 Split；与人工选择结果对比。",
        "按 selfcapability / 机台组小范围启用，验证 Split、R2R 回执与对账。",
        "逐步扩大范围，并持续观察 over Q-time、效率、失败率与重复选择指标。",
        "达到稳定门槛后转为全量运行，保留人工暂停开关。",
    ]:
        add_list_item(doc, item, rollout_decimal_id)
    add_heading(doc, "回退策略", 2)
    add_body(doc, "优先通过 AMA.TriggerConfig 将 Switch 置为非 Y 停止新任务；对已 Split 但未完成 R2R 同步的记录进入人工补偿清单。回退不得直接删除审计记录或绕过母子批一致性检查。")

    add_heading(doc, "待确认事项", 1)
    add_table(doc, ["编号", "事项", "建议责任方", "关闭条件"], [
        ["O-01", "可用 Lot 状态白名单及 qty>5 是否适用于所有 CMP 场景。", "CMP 工艺 / CMPAssign", "形成状态与数量规则表。"],
        ["O-02", "Q-time 风险 Lot 的 Remaining 排序、空值及“无 Qtimelot”处理。", "CMP 工艺", "用示例数据确认排序结果。"],
        ["O-03", "LimitWIP 公式各字段单位、Min(QTimeLimit) 取值范围与 Ratio 精度。", "业务 / RTD", "通过边界测试与样例计算。"],
        ["O-04", "Reasoncode_CMP 允许/排除集合及 LowWipControl 的确切移除时点。", "PE / R2R", "形成接口枚举与时序图。"],
        ["O-05", "CMPLowWipMachinedouble 的准确参数名、多值格式与特殊边清单。", "RTD / 运维", "配置字典评审通过。"],
        ["O-06", "Split 与 R2R 的接口字段、幂等键、超时、重试与补偿协议。", "AMA / R2R", "接口契约签字确认。"],
        ["O-07", "冷却期 1 天的起算时点、解除条件与跨日批次处理。", "业务 / 开发", "验收用例覆盖。"],
    ], [800, 4500, 1800, 2260], font_size=8.55)

    add_heading(doc, "附录：伪代码", 1)
    add_code_block(doc, [
        "on_watchdog_tick():",
        "  if not trigger_config.enabled(now): return",
        "  with non_reentrant_lock('CMPAutoPirun'):",
        "    snapshot = load_latest_cmpassign_and_configs()",
        "    groups = group_by_selfcapability(snapshot)",
        "    for group in groups:",
        "      while group.selected_count < group.picount_limit:",
        "        lot = select_risk_lot(group) or select_highwip_lot(group)",
        "        if lot is None: break",
        "        realtime_lot = double_check(lot)",
        "        tool = select_low_loading_safe_tool(realtime_lot)",
        "        if tool is None: mark_unassigned(lot); continue",
        "        child = idempotent_split(lot, configured_split_qty)",
        "        notify_r2r(child, tool, merge_info)",
        "        write_audit_chain(snapshot, lot, tool, child)",
    ])
    add_heading(doc, "附录：来源图片清单", 2)
    add_body(doc, "D:\\Obsidian\\work\\OBSidianCodex\\00.raw-materials\\10.sources\\images\\CMPAutoPirun", italic=True)
    for item in ["技术文档1.jpg - 技术文档4.jpg", "需求单1.jpg - 需求单5.jpg", "Testcase1.jpg - Testcase2.jpg", "PPT.jpg"]:
        add_list_item(doc, item, bullet_id)

    doc.core_properties.title = "CMPAutoPirun 详细设计说明书"
    doc.core_properties.subject = "CMP R2R Pi-Run Lot 自动分批功能详细设计"
    doc.core_properties.author = "CMPAutoPirun 项目组"
    doc.core_properties.keywords = "CMP, R2R, Pi-Run, AMA, 自动分批, 详细设计"
    doc.core_properties.comments = "Based on project source screenshots; review draft."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
